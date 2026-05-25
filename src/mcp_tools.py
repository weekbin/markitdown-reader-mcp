# -*- coding: utf-8 -*-
"""MCP tool functions and business logic extracted from server.py"""

from __future__ import annotations

import json
import logging
import os
import resource
import traceback
from pathlib import Path
from typing import Optional

from mcp.server.fastmcp import FastMCP

from .constants import BASE_DIR, IMAGE_SIZE_THRESHOLD, SLICE_BLOCKS, SLICE_PAGES
from .storage import (
    _calculate_content_hash,
    _get_doc_dir,
    _get_index_path,
    _get_slices_dir,
    _load_index,
    _lock_file,
    _make_doc_name,
    _move_current_to_history,
    _save_index,
    _save_index_nolock,
    _unlock_file,
    _with_read_lock,
)
from .parser import (
    _extract_images_from_docx,
    _extract_images_from_pdf,
    _get_docx_image_anchor,
    _is_pdf_by_magic,
    _read_docx_text,
    _read_generic_text,
    _read_pdf_text,
    _slice_docx,
    _slice_pdf,
)
from .image import _is_small_image, _ocr_small_image
from .callbacks import _post_callback
from .utils import mem

_log = logging.getLogger("markitdown")

mcp = FastMCP("markitdown-reader")


# ─────────────────────────────────────────────────────────────────
# MCP 工具
# ─────────────────────────────────────────────────────────────────

@mcp.tool()
def read_document(file_path: str, fast: bool = False, slice_ids: Optional[list[str]] = None, force_refresh: bool = False, callback_url: str = "") -> str:
    _log.debug(f"read_document CALLED file={file_path} fast={fast} slice_ids={slice_ids} force_refresh={force_refresh} callback_url={callback_url} mem={mem()}MB")
    if not os.path.isfile(file_path):
        return f"[错误] 文件不存在: {file_path}"

    doc_name = _make_doc_name(file_path)

    # force_refresh: 清理旧缓存
    if force_refresh:
        _move_current_to_history(doc_name)

    paired = _find_paired_file(file_path)
    extract_imgs = not fast

    try:
        # 优先判断配对：有配对文件时，始终走 _read_paired_documents
        #（即使提供了 slice_ids，也通过配对路径读取，以保证图片提取和 index 类型正确）
        if paired is not None:
            ext = Path(file_path).suffix.lower()
            if ext == ".pdf":
                result = _read_paired_documents(file_path, paired, extract_images=extract_imgs, callback_url=callback_url, force_refresh=force_refresh)
            else:
                result = _read_paired_documents(paired, file_path, extract_images=extract_imgs, callback_url=callback_url, force_refresh=force_refresh)
        elif slice_ids is not None:
            ext = Path(file_path).suffix.lower()
            is_pdf = ext == ".pdf" or _is_pdf_by_magic(file_path)
            result = _read_slices_direct(doc_name, slice_ids, is_pdf, extract_imgs, force_refresh=force_refresh)
            result["paired"] = False
        else:
            result = _read_single_document(file_path, doc_name, extract_images=extract_imgs, callback_url=callback_url, force_refresh=force_refresh)
    except Exception:
        _log.exception(f"  read_document CRASHED: {traceback.format_exc()}")
        return f"[错误] 处理失败:\n{traceback.format_exc()}"

    # 构造返回
    img_lines = []
    ocr_count = 0
    for img in result["images"]:
        size_kb = img["size"] // 1024
        ocr_hint = f" [OCR: {img.get('ocr', '')}]" if img.get('ocr') else ""
        img_lines.append(f"[图片: file://{img['path']}] ({size_kb}KB){ocr_hint}")
        if img.get('ocr'):
            ocr_count += 1

    # 配对模式下同时显示 PDF 和 DOCX 分片信息
    if result["paired"]:
        pdf_slices = result.get("pdf_slices", [])
        docx_slices = result.get("docx_slices", result["slices"])
        pdf_info = ", ".join(s["id"] for s in pdf_slices)
        docx_info = ", ".join(s["id"] for s in docx_slices)
        slice_info = f"PDF {len(pdf_slices)}片 ({pdf_info}) | DOCX {len(docx_slices)}片 ({docx_info})"
    else:
        slice_info = ", ".join(s["id"] for s in result["slices"])
    img_summary = f"{len(result['images'])}张图片"
    if ocr_count > 0:
        img_summary += f"（含{ocr_count}张已OCR）"

    paired_line = (
        f"配对: {'是 → ' + Path(result['paired_file']).name if result['paired'] else '未找到配对文件'}\n"
    )
    if not result["paired"]:
        ext = Path(file_path).suffix.lower()
        if ext == ".pdf":
            paired_line += (
                "⚠️ 只提供 PDF: 文字可能乱码（CID-font），图片可正常提取\n"
                "提示: 如有 DOCX 版本，建议 read_document_pair(pdf, docx)\n"
                + ("⚠️ fast=True 已跳过图片提取，图片请用 extract_images() 单独提取\n" if fast else "      继续干活...\n")
            )
        else:
            paired_line += (
                "⚠️ 只提供 DOCX: 文字表格可正常提取，可能无法提取图片\n"
                "提示: 如有 PDF 版本，建议 read_document_pair(pdf, docx)\n"
                + ("⚠️ fast=True 已跳过图片提取，图片请用 extract_images() 单独提取\n" if fast else "      继续干活...\n")
            )

    # 结构化状态 JSON（agent 可直接解析）
    pdf_slices_status = result.get("pdf_slices", [])
    docx_slices_status = result.get("docx_slices", result["slices"])
    
    next_steps = []
    if result["paired"] and not result["written"]:
        next_steps.append("Use read_document_pair() for optimal PDF+DOCX processing")
    if not result["paired"] and result["images"]:
        next_steps.append(f"Consider using extract_images() to re-extract with different settings")
    if ocr_count == 0 and result["images"]:
        next_steps.append(f"Small images may benefit from OCR - try ocr_image() on specific images")
    if result["written"]:
        next_steps.append(f"Content written to {result['content_path']} - read for full text")
    if result["slices"]:
        next_steps.append(f"Use slice_ids in read_document() to re-read specific slices")
    
    status = {
        "doc": Path(file_path).name,
        "paired": result["paired"],
        "paired_file": Path(result["paired_file"]).name if result["paired_file"] else None,
        "need_pairing": not result["paired"],
        "fast": fast,
        "images_extracted": extract_imgs,
        "slice_count": len(result["slices"]),  # 兼容：PDF 单文档时为 PDF 片数，配对时为 DOCX 片数
        "pdf_slice_count": len(pdf_slices_status),
        "docx_slice_count": len(docx_slices_status),
        "slices": [s["id"] for s in result["slices"]],
        "pdf_slices": [s["id"] for s in pdf_slices_status],
        "docx_slices": [s["id"] for s in docx_slices_status],
        "image_count": len(result["images"]),
        "ocr_count": ocr_count,
        "content_written": result["written"],
        "content_path": result["content_path"],
        "chars": len(result["text"]),
        "mode": "paired" if result["paired"] else ("pdf_single" if (Path(file_path).suffix.lower() == ".pdf") else "docx_single"),
        "slicing_mode": "paired" if result["paired"] else ("pdf_only" if (Path(file_path).suffix.lower() == ".pdf") else "docx_only"),
        "next_steps": next_steps,
    }
    status_json = json.dumps(status, ensure_ascii=False)

    progress_msgs = result.get("progress", [])
    progress_line = ""
    if progress_msgs:
        progress_line = "⚙️ " + " | ".join(progress_msgs) + "\n"

    header = (
        f"[STATUS]\n{status_json}\n[/STATUS]\n\n"
        + progress_line
        + f"# {Path(file_path).name}\n"
        + paired_line
        + f"分片: {len(result['slices'])}片 ({slice_info})\n"
        + f"图片: {img_summary}\n"
    )

    if result["written"]:
        footer = (
            f"\n---\n[内容已写入文件: {result['content_path']}]\n"
            f"[完整内容请读取该文件]"
        )
    else:
        footer = f"\n--- 字符: {len(result['text'])} ---\n"

    if img_lines:
        footer += "\n" + "\n".join(img_lines)

    if force_refresh:
        index = _load_index(doc_name)
        index["content_hash"] = _calculate_content_hash(file_path)
        index["source_file"] = file_path
        _save_index(doc_name, index)

    return header + footer


@mcp.tool()
def read_document_pair(pdf_path: str, docx_path: str, force_refresh: bool = False, callback_url: str = "") -> str:
    """显式配对读取：PDF 提取图片，DOCX 提取文字和表格"""
    _log.debug(f"read_document_pair CALLED pdf={pdf_path} docx={docx_path} callback_url={callback_url} mem={mem()}MB")
    if not os.path.isfile(pdf_path):
        return f"[错误] PDF 不存在: {pdf_path}"
    if not os.path.isfile(docx_path):
        return f"[错误] DOCX 不存在: {docx_path}"

    doc_name = _make_doc_name(pdf_path)
    if force_refresh:
        _move_current_to_history(doc_name)
    try:
        result = _read_paired_documents(pdf_path, docx_path, callback_url=callback_url, force_refresh=force_refresh)
    except Exception:
        _log.exception(f"  read_document_pair CRASHED: {traceback.format_exc()}")
        return f"[错误] 处理失败:\n{traceback.format_exc()}"

    img_lines = []
    ocr_count = 0
    for img in result["images"]:
        size_kb = img["size"] // 1024
        ocr_hint = f" [OCR: {img.get('ocr', '')}]" if img.get('ocr') else ""
        img_lines.append(f"[图片: file://{img['path']}] ({size_kb}KB){ocr_hint}")
        if img.get('ocr'):
            ocr_count += 1

    pdf_info = ", ".join(s["id"] for s in result["pdf_slices"])
    docx_info = ", ".join(s["id"] for s in result["docx_slices"])
    img_summary = f"{len(result['images'])}张图片"
    if ocr_count > 0:
        img_summary += f"（含{ocr_count}张已OCR）"

    header = (
        f"# {Path(pdf_path).name} + {Path(docx_path).name}\n"
        f"配对: 是（PDF图片 + DOCX文字表格）\n"
        f"分片: PDF {len(result['pdf_slices'])}片 ({pdf_info}) | DOCX {len(result['docx_slices'])}片 ({docx_info})\n"
        f"图片: {img_summary}\n"
    )

    if result["written"]:
        footer = (
            f"\n---\n[内容已写入文件: {result['content_path']}]\n"
            f"[完整内容请读取该文件]"
        )
    else:
        footer = f"\n--- 字符: {len(result['text'])} ---\n"

    if img_lines:
        footer += "\n" + "\n".join(img_lines)

    if force_refresh:
        index = _load_index(doc_name)
        index["content_hash"] = _calculate_content_hash(pdf_path)
        index["source_file"] = pdf_path
        _save_index(doc_name, index)

    next_steps = []
    if ocr_count == 0 and result["images"]:
        next_steps.append("Small images may benefit from OCR - try ocr_image() on specific images")
    if result["written"]:
        next_steps.append(f"Content written to {result['content_path']} - read for full text")
    if result["pdf_slices"]:
        next_steps.append(f"Use slice_document() to see PDF page distribution")

    if next_steps:
        footer += "\n\n下一步建议:\n" + "\n".join(f"- {s}" for s in next_steps)

    return header + footer


@mcp.tool()
def extract_images(file_path: str, page_range: str = "") -> str:
    """
    专门提取图片。返回图片列表。
    page_range: 可选，如 "1-5" 或 "1,3,5"
    """
    import tempfile
    if not os.path.isfile(file_path):
        return f"[错误] 文件不存在: {file_path}"

    doc_name = _make_doc_name(file_path)
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf" or _is_pdf_by_magic(file_path):
        if page_range:
            import fitz
            with fitz.open(file_path) as doc:
                indices = []
                for part in page_range.split(","):
                    part = part.strip()
                    if "-" in part:
                        s, e = part.split("-", 1)
                        indices.extend(range(int(s) - 1, int(e)))
                    else:
                        indices.append(int(part) - 1)
                with fitz.open() as tmp:
                    for idx in indices:
                        if 0 <= idx < len(doc):
                            tmp.insert_pdf(doc, from_page=idx, to_page=idx)
                    tmp_path = tempfile.mktemp(suffix=".pdf")
                    tmp.save(tmp_path)
            imgs, index = _extract_images_from_pdf(tmp_path, doc_name + "_range")
            _save_index(doc_name + "_range", index)
            os.unlink(tmp_path)
        else:
            imgs, index = _extract_images_from_pdf(file_path, doc_name)
            _save_index(doc_name, index)
    else:
        imgs = _extract_images_from_docx(file_path, doc_name)

    lines = [f"共 {len(imgs)} 张图片:\n"]
    for img in imgs:
        size_kb = img["size"] // 1024
        ocr_hint = f" [OCR] {img.get('ocr', '')}" if img.get('ocr') else ""
        lines.append(f"[图片: file://{img['path']}] ({size_kb}KB){ocr_hint}")

    next_steps = [
        "Use ocr_image() for individual image OCR if not auto-processed",
        "Use read_document() to extract text alongside images",
    ]
    if imgs:
        next_steps.append(f"Images stored in cache - use get_cached_content('{doc_name}') to retrieve")

    lines.append("\n下一步建议:")
    for s in next_steps:
        lines.append(f"- {s}")

    return "\n".join(lines)


@mcp.tool()
def ocr_image(image_path: str) -> str:
    """
    对单张图片进行 OCR 识别（内置 tesseract）。
    支持 PNG/JPG/BMP 等常见格式。
    返回识别文字，失败返回空字符串。
    """
    if not os.path.isfile(image_path):
        return f"[错误] 文件不存在: {image_path}"
    _log.debug(f"ocr_image CALLED path={image_path} mem={mem()}MB")
    text = _ocr_small_image(image_path)
    if text:
        _log.debug(f"ocr_image: OK, {len(text)} chars")
        result = text + "\n\n[next_steps]\nUse update_document_markdown() to save OCR result to document index"
    else:
        _log.debug(f"ocr_image: no text found")
        result = "[next_steps]\nNo text found - image may be empty or need manual transcription"
    return result


@mcp.tool()
def slice_document(file_path: str, pages_per_slice: int = 5) -> str:
    """
    分片工具：只做分片，返回分片信息（不读取内容）。
    用于提前了解分片结构，或单独控制分片逻辑。
    返回: JSON 格式分片列表
    """
    _log.debug(f"slice_document CALLED file={file_path} mem={mem()}MB")
    if not os.path.isfile(file_path):
        return f'[{{"error": "文件不存在: {file_path}"}}]'

    doc_name = _make_doc_name(file_path)
    ext = Path(file_path).suffix.lower()

    try:
        if ext == ".pdf" or _is_pdf_by_magic(file_path):
            import fitz
            with fitz.open(file_path) as doc:
                total = len(doc)
                slices = []
                for i in range(0, total, pages_per_slice):
                    end = min(i + pages_per_slice, total)
                    slices.append({
                        "slice_id": f"p{i+1}-{end}",
                        "page_range": f"{i+1}-{end}",
                        "page_count": end - i,
                        "slice_index": i // pages_per_slice,
                    })
                result_json = {
                    "doc_name": doc_name,
                    "total_pages": total,
                    "slice_count": len(slices),
                    "slices": slices,
                    "paired_file": _find_paired_file(file_path),
                    "next_steps": ["Use read_document() with slice_ids to read specific slices", "Use extract_images() to extract images from the document"],
                }
            return json.dumps(result_json, ensure_ascii=False, indent=2)

        elif ext in (".docx", ".doc"):
            from docx import Document
            doc = Document(file_path)
            blocks = sum(1 for _ in doc.element.body)
            slices = []
            for i in range(0, blocks, SLICE_BLOCKS):
                end = min(i + SLICE_BLOCKS, blocks)
                slices.append({
                    "slice_id": f"b{i+1}-{end}",
                    "block_range": f"{i+1}-{end}",
                    "block_count": end - i,
                    "slice_index": i // SLICE_BLOCKS,
                })
            result_json = {
                "doc_name": doc_name,
                "total_blocks": blocks,
                "slice_count": len(slices),
                "slices": slices,
                "paired_file": _find_paired_file(file_path),
                "next_steps": ["Use read_document() with slice_ids to read specific slices", "Use extract_images() to extract images from the document"],
            }
            return json.dumps(result_json, ensure_ascii=False, indent=2)

        else:
            return f'[{{"error": "不支持分片格式: {ext}"}}]'

    except Exception as e:
        _log.exception(f"  slice_document CRASHED for {file_path}: {traceback.format_exc()}")
        return f'[{{"error": "{type(e).__name__}: {e}"}}]'


@mcp.tool()
def get_document_info(file_path: str) -> str:
    """返回文档元信息：页数、分片数、图片数、配对建议"""
    _log.debug(f"get_document_info CALLED file={file_path} mem={mem()}MB")
    if not os.path.isfile(file_path):
        return f"[错误] 文件不存在: {file_path}"

    doc_name = _make_doc_name(file_path)
    size = os.path.getsize(file_path)
    ext = Path(file_path).suffix.lower()

    info = [
        f"文件: {Path(file_path).name}",
        f"大小: {size / 1024:.0f} KB ({size / 1024**2:.1f} MB)",
    ]

    paired = _find_paired_file(file_path)
    if paired:
        info.append(f"配对文件: 有 → {Path(paired).name}")
        info.append("建议: 使用 read_document_pair() 显式配对读取")
    else:
        info.append(f"配对文件: 无")
        info.append("提示: 如有另一格式版本，建议提供以获得更完整结果")

    if ext == ".pdf" or _is_pdf_by_magic(file_path):
        try:
            import fitz
            with fitz.open(file_path) as doc:
                total = len(doc)
                slices = (total + SLICE_PAGES - 1) // SLICE_PAGES
                info.append(f"页数: {total}")
                info.append(f"预估分片数: {slices}片 (每{SLICE_PAGES}页)")
        except Exception as e:
            _log.warning(f"  get_document_info: fitz failed for {file_path}: {e}")
    elif ext in (".docx", ".doc"):
        try:
            from docx import Document
            doc = Document(file_path)
            blocks = sum(1 for _ in doc.element.body)
            slices = (blocks + SLICE_BLOCKS - 1) // SLICE_BLOCKS
            info.append(f"段落: {len(doc.paragraphs)}")
            info.append(f"表格: {len(doc.tables)}")
            info.append(f"预估分片数: {slices}片 (每{SLICE_BLOCKS}block)")
        except Exception as e:
            _log.warning(f"  get_document_info: docx failed for {file_path}: {e}")

    # 检查图片池
    index = _load_index(doc_name)
    if index["images"]:
        info.append(f"图片池已有: {len(index['images'])} 张")

    next_steps = [
        "Use read_document() to extract full content",
        "Use slice_document() to preview slice structure",
        "Use extract_images() to extract images separately",
    ]
    if paired:
        next_steps.insert(0, "Use read_document_pair() for optimal PDF+DOCX processing")

    info.append("\n下一步建议:")
    for s in next_steps:
        info.append(f"- {s}")

    return "\n".join(info)


@mcp.tool()
def list_supported_files() -> str:
    return (
        "支持的文件格式:\n"
        "  PDF:  .pdf\n"
        "  Word: .docx, .doc\n"
        "  文本: .txt, .md, .json, .xml, .csv, .html\n"
        "\n"
        "工具接口:\n"
        "  read_document(path)         - 主工具，自动配对检测\n"
        "  read_document_pair(pdf, docx) - 显式配对读取\n"
        "  extract_images(path)        - 专门提取图片\n"
        "  get_document_info(path)      - 文档元信息\n"
        "\n"
        "下一步建议:\n"
        "- Use read_document() to start extracting document content\n"
        "- Use get_document_info() to check document structure first\n"
    )


@mcp.tool()
def list_cache_dir(doc_name: str = None) -> str:
    """
    列出缓存目录。
    - doc_name: 可选，指定文档名称过滤
    不指定则返回所有已缓存文档列表。
    """
    if doc_name:
        doc_dir = _get_doc_dir(doc_name)
        if not doc_dir.exists():
            return json.dumps({"error": f"Document {doc_name} not found in cache"}, ensure_ascii=False)
        index = _load_index(doc_name)
        result = {
            "doc_name": doc_name,
            "path": str(doc_dir),
            "exists": True,
            "has_slices": (doc_dir / "slices").exists(),
            "has_images": (doc_dir / "images").exists(),
            "slice_count": len(index.get("pdf_slices", [])) + len(index.get("docx_slices", [])),
            "image_count": len(index.get("images", [])),
            "next_steps": ["Use get_cached_content() to read cached content", "Use get_processing_status() to check processing state"],
        }
        return json.dumps(result, ensure_ascii=False, indent=2)
    else:
        docs = []
        for d in BASE_DIR.iterdir():
            if not d.is_dir():
                continue
            index_path = d / "index.json"
            if not index_path.exists():
                continue
            index = _load_index(d.name)
            docs.append({
                "doc_name": d.name,
                "path": str(d),
                "slice_count": len(index.get("pdf_slices", [])) + len(index.get("docx_slices", [])),
                "image_count": len(index.get("images", [])),
                "source_file": index.get("source_file", ""),
            })
        result = {
            "cached_documents": docs,
            "total": len(docs),
            "next_steps": ["Use list_cache_dir(doc_name='xxx') for specific document details", "Use get_cached_content() to read cached content"],
        }
        return json.dumps(result, ensure_ascii=False, indent=2)


@mcp.tool()
def update_document_markdown(img_id: str, ocr_result: str, position_info: dict = None) -> str:
    """
    更新单张图片的OCR结果到index.json和output.md。
    - img_id: 图片文件名或MD5
    - ocr_result: OCR识别文字
    - position_info: 可选，位置信息字典
    """
    _log.debug(f"update_document_markdown CALLED img_id={img_id}")

    for img_dir in BASE_DIR.iterdir():
        if not img_dir.is_dir():
            continue
        index_path = img_dir / "index.json"
        if not index_path.exists():
            continue
        index = _load_index(img_dir.name)
        for img in index.get("images", []):
            if img.get("name") == img_id or img.get("md5") == img_id:
                doc_name = img_dir.name
                if _update_image_ocr(doc_name, img_id, ocr_result, position_info or {}):
                    _rebuild_output_md(doc_name)
                    return json.dumps({
                        "updated": True,
                        "img_id": img_id,
                        "doc_name": doc_name,
                        "ocr_result": ocr_result[:100] + "..." if len(ocr_result) > 100 else ocr_result,
                        "next_steps": ["Use get_cached_content() to read updated document", "Use get_processing_status() to check overall progress"],
                    }, ensure_ascii=False)
                return json.dumps({"updated": False, "error": "Failed to update", "next_steps": ["Check if image ID is correct", "Verify document cache exists"]}, ensure_ascii=False)

    return json.dumps({"updated": False, "error": f"Image {img_id} not found in any document", "next_steps": ["Verify the image was extracted", "Check if document was processed"]}, ensure_ascii=False)


@mcp.tool()
def update_batch_document_markdown(updates: list[dict]) -> str:
    """
    批量更新多张图片的OCR结果。
    updates: [{"img_id": str, "ocr_result": str, "position_info": dict}, ...]
    返回: {"updated": int, "remaining": int, "failed": list}
    """
    _log.debug(f"update_batch_document_markdown CALLED with {len(updates)} updates")

    updated = 0
    remaining = 0
    failed = []

    updated_docs = set()

    for update in updates:
        img_id = update.get("img_id", "")
        ocr_result = update.get("ocr_result", "")
        position_info = update.get("position_info") or {}

        found = False
        for img_dir in BASE_DIR.iterdir():
            if not img_dir.is_dir():
                continue
            index_path = img_dir / "index.json"
            if not index_path.exists():
                continue
            index = _load_index(img_dir.name)
            for img in index.get("images", []):
                if img.get("name") == img_id or img.get("md5") == img_id:
                    doc_name = img_dir.name
                    if _update_image_ocr(doc_name, img_id, ocr_result, position_info):
                        updated += 1
                        found = True
                        updated_docs.add(doc_name)
                        break
            if found:
                break

        if not found:
            remaining += 1
            failed.append(img_id)

    for doc_name in updated_docs:
        _rebuild_output_md(doc_name)

    return json.dumps({
        "updated": updated,
        "remaining": remaining,
        "failed": failed,
        "next_steps": ["Use get_cached_content() to read updated documents", "Use get_processing_status() to check overall progress", f"Retry {remaining} failed images individually if needed"],
    }, ensure_ascii=False)


@mcp.tool()
def get_processing_status(doc_name: str = None, file_path: str = None) -> str:
    """
    获取文档处理状态。
    - doc_name: 文档名称（可选，与file_path二选一）
    - file_path: 文件路径（可选，会自动转换为doc_name）
    返回: {pending_images, completed_images, failed_images, progress%, next_steps}
    """
    if file_path:
        doc_name = _make_doc_name(file_path)

    if not doc_name:
        return json.dumps({"error": "doc_name or file_path required"}, ensure_ascii=False)

    def _get_status():
        index = _load_index(doc_name)
        images = index.get("images", [])

        pending_images = []
        completed_images = []
        failed_images = []

        for img in images:
            status = img.get("ocr_status", "pending")
            img_info = {
                "name": img.get("name"),
                "page": img.get("page"),
                "size": img.get("size"),
            }
            if status == "completed":
                completed_images.append(img_info)
            elif status == "failed":
                failed_images.append(img_info)
            else:
                pending_images.append(img_info)

        total = len(images)
        progress = (len(completed_images) / total * 100) if total > 0 else 0

        next_steps = []
        if pending_images:
            next_steps.append(f"Retry {len(pending_images)} pending images with retry_failed_images()")
        if failed_images:
            next_steps.append(f"Review {len(failed_images)} failed images - they may need manual OCR")
        if not next_steps:
            next_steps.append("All images processed - document is complete")

        slicing_mode = "paired" if index.get("pdf_slices") and index.get("docx_slices") else ("pdf_only" if index.get("pdf_slices") else "docx_only" if index.get("docx_slices") else "single")

        return {
            "doc_name": doc_name,
            "pending_images": pending_images,
            "completed_images": completed_images,
            "failed_images": failed_images,
            "progress": round(progress, 1),
            "total_images": total,
            "next_steps": next_steps,
            "slicing_mode": slicing_mode,
        }

    try:
        result = _with_read_lock(doc_name, _get_status)
        return json.dumps(result, ensure_ascii=False, indent=2)
    except TimeoutError as e:
        return json.dumps({"error": str(e)}, ensure_ascii=False)


@mcp.tool()
def retry_failed_images(doc_name: str) -> str:
    """
    重试失败图片的OCR处理。
    - doc_name: 文档名称
    返回: {retried: int, succeeded: int, still_failing: int}
    """
    _log.debug(f"retry_failed_images CALLED doc_name={doc_name}")
    index = _load_index(doc_name)

    retried = 0
    succeeded = 0
    still_failing = []

    for img in index.get("images", []):
        if img.get("ocr_status") == "failed":
            img_path = img.get("path")
            if img_path and os.path.isfile(img_path):
                retried += 1
                ocr_text = _ocr_small_image(img_path)
                if ocr_text:
                    img["ocr_status"] = "completed"
                    img["ocr_result"] = ocr_text
                    img["ocr_source"] = "tesseract_retry"
                    succeeded += 1
                else:
                    still_failing.append(img.get("name"))

    _save_index(doc_name, index)
    _rebuild_output_md(doc_name)

    return json.dumps({
        "retried": retried,
        "succeeded": succeeded,
        "still_failing": len(still_failing),
        "failed_names": still_failing,
        "next_steps": ["Use get_cached_content() to read document with OCR results", "Use update_document_markdown() to manually update still-failing images", "Use get_processing_status() to check overall progress"],
    }, ensure_ascii=False)


@mcp.tool()
def resume_document(file_path: str) -> str:
    """
    恢复不完整的文档处理。
    - file_path: 文档路径
    自动从index.json检测未完成的处理
    返回状态类似read_document
    """
    _log.debug(f"resume_document CALLED file={file_path}")
    if not os.path.isfile(file_path):
        return f"[错误] 文件不存在: {file_path}"

    doc_name = _make_doc_name(file_path)
    index = _load_index(doc_name)

    if not index.get("images") and not index.get("pdf_slices") and not index.get("docx_slices"):
        return json.dumps({
            "error": "No processing history found. Use read_document() first.",
            "doc_name": doc_name,
        }, ensure_ascii=False)

    pending_imgs = [img for img in index.get("images", []) if img.get("ocr_status") != "completed"]
    failed_imgs = [img for img in index.get("images", []) if img.get("ocr_status") == "failed"]
    completed_imgs = [img for img in index.get("images", []) if img.get("ocr_status") == "completed"]

    slicing_mode = "paired" if index.get("pdf_slices") and index.get("docx_slices") else ("pdf_only" if index.get("pdf_slices") else "docx_only" if index.get("docx_slices") else "single")

    pdf_slices = index.get("pdf_slices", [])
    docx_slices = index.get("docx_slices", [])
    slices = docx_slices if docx_slices else pdf_slices

    doc_dir = _get_doc_dir(doc_name)
    output_path = doc_dir / "output.md"
    output_exists = output_path.exists()

    result = {
        "doc_name": doc_name,
        "slicing_mode": slicing_mode,
        "resume": True,
        "pdf_slice_count": len(pdf_slices),
        "docx_slice_count": len(docx_slices),
        "total_images": len(index.get("images", [])),
        "pending_images": len(pending_imgs),
        "failed_images": len(failed_imgs),
        "completed_images": len(completed_imgs),
        "progress": round(len(completed_imgs) / len(index.get("images", [])) * 100, 1) if index.get("images") else 0,
        "output_exists": output_exists,
        "output_path": str(output_path) if output_exists else None,
        "pdf_slices": [s["id"] for s in pdf_slices],
        "docx_slices": [s["id"] for s in docx_slices],
        "slices": [s["id"] for s in slices],
        "next_steps": [],
    }

    if pending_imgs:
        result["next_steps"].append(f"Process {len(pending_imgs)} pending images")
    if failed_imgs:
        result["next_steps"].append(f"Retry {len(failed_imgs)} failed images with retry_failed_images('{doc_name}')")
    if not result["next_steps"]:
        result["next_steps"].append("Document processing complete")

    return json.dumps(result, ensure_ascii=False, indent=2)


@mcp.tool()
def get_cached_content(doc_name: str, run: str = "latest") -> str:
    """
    获取缓存文档的内容。
    - doc_name: 文档名称
    - run: "latest"（默认）或 "run_XXX" 格式的历史run
    返回: 文档内容和元信息
    """
    _log.debug(f"get_cached_content CALLED doc_name={doc_name} run={run} mem={mem()}MB")

    doc_dir = _get_doc_dir(doc_name)

    if run == "latest":
        content_path = doc_dir / "output.md"
        index = _load_index(doc_name)
    else:
        run_dir = doc_dir / "history" / run
        if not run_dir.exists():
            return json.dumps({"error": f"Run {run} not found for {doc_name}"}, ensure_ascii=False)
        content_path = run_dir / "output.md"
        index_path = run_dir / "index.json"
        if index_path.exists():
            try:
                index = json.loads(index_path.read_text(encoding="utf-8"))
            except Exception:
                index = {}
        else:
            index = {}

    if not content_path.exists():
        return json.dumps({"error": f"No content found for {doc_name}", "doc_name": doc_name}, ensure_ascii=False)

    try:
        content = content_path.read_text(encoding="utf-8")
    except Exception as e:
        return json.dumps({"error": f"Failed to read content: {e}"}, ensure_ascii=False)

    result = {
        "doc_name": doc_name,
        "run": run,
        "content": content,
        "chars": len(content),
        "index": {
            "slices": index.get("pdf_slices", []) + index.get("docx_slices", []),
            "images": len(index.get("images", [])),
            "source_file": index.get("source_file", ""),
            "content_hash": index.get("content_hash", ""),
        },
        "next_steps": ["Use extract_images() to re-extract images if needed", "Use get_processing_status() to check processing state"],
    }
    return json.dumps(result, ensure_ascii=False, indent=2)


# ─────────────────────────────────────────────────────────────────
# Business logic functions
# ─────────────────────────────────────────────────────────────────

def _find_paired_file(path: str) -> Optional[str]:
    """
    在同目录查找同名不同扩展名的配对文件。

    跳过已经是分片的文件（path 包含 /slices/ 或文件名含 slice_ 前缀），
    避免对分片文件错误触发配对逻辑导致重复分片和无限递归。
    """
    p = Path(path)
    if "/slices/" in str(p) or p.name.startswith("slice_"):
        return None
    stem = p.stem
    directory = p.parent
    for ext in [".pdf", ".docx"]:
        if p.suffix.lower() != ext:
            paired = directory / f"{stem}{ext}"
            if paired.exists():
                return str(paired)
    return None


def _read_slices_direct(
    doc_name: str,
    slice_ids: list[str],
    is_pdf: bool,
    extract_images: bool = True,
    force_refresh: bool = False,
) -> dict:
    """
    直接读取预存的分片文件，不重新切片。
    slice_ids 格式：["b1-200", "b201-400", ...]（DOCX）或 ["p1-5", "p6-10", ...]（PDF）
    分片文件命名：slice_000.docx, slice_001.pdf 等，按顺序索引。
    index 使用 pdf_slices / docx_slices 分开存储（兼容旧 slices 格式）。
    当 index 条目少于实际 slice 文件时，从实际文件读取页码/块范围。
    """
    _log.debug(f"  _read_slices_direct: doc={doc_name} slices={slice_ids} is_pdf={is_pdf} mem={mem()}MB")
    slices_dir = _get_slices_dir(doc_name)

    slice_files = sorted(slices_dir.glob("slice_*.docx" if not is_pdf else "slice_*.pdf"))
    index = _load_index(doc_name)
    # 优先使用分类型 index（pdf_slices / docx_slices），兼容旧格式（slices）
    slices_key = "pdf_slices" if is_pdf else "docx_slices"
    stored = index.get(slices_key, index.get("slices", []))
    stored_ids = [s["id"] for s in stored]

    def _pdf_page_range(pdf_path: str) -> tuple[int, int]:
        """从 PDF slice 文件读取实际页码范围。"""
        import fitz
        with fitz.open(pdf_path) as d:
            total = len(d)
            return 1, total  # slice 文件只含连续页，起止页可从文件名推导

    id_to_path: dict[str, str] = {}
    path_to_index: dict[str, int] = {}
    for i, sf in enumerate(slice_files):
        path_str = str(sf)
        path_to_index[path_str] = i
        if i < len(stored_ids):
            id_to_path[stored_ids[i]] = path_str
        else:
            # index 条目不足，从文件名推导 slice index，再用 SLICE_PAGES 计算实际范围
            if is_pdf:
                # slice_037.pdf → i=37, pages 186-187 (假设 187 总页)
                # 先尝试从 index 推断总页数
                total_pages = stored[0]["total"] if stored else i * SLICE_PAGES + SLICE_PAGES
                start = i * SLICE_PAGES + 1
                end = min((i + 1) * SLICE_PAGES, total_pages)
                id_to_path[f"p{start}-{end}"] = path_str
            else:
                total_blocks = stored[0]["total"] if stored else i * SLICE_BLOCKS + SLICE_BLOCKS
                start = i * SLICE_BLOCKS + 1
                end = min((i + 1) * SLICE_BLOCKS, total_blocks)
                id_to_path[f"b{start}-{end}"] = path_str

    # 只读请求的分片
    text_parts = []
    found_ids = []
    all_images = []

    for sid in slice_ids:
        if sid not in id_to_path:
            _log.warning(f"  _read_slices_direct: slice {sid} not found, skipping")
            continue
        path = id_to_path[sid]
        if is_pdf:
            txt = _read_pdf_text(path)
        else:
            txt = _read_docx_text(path)
        text_parts.append(f"=== [{sid}] ===\n{txt}")
        found_ids.append(sid)
        if extract_images:
            if is_pdf:
                slice_idx = path_to_index.get(path, 0)
                starting_page = slice_idx * SLICE_PAGES + 1
                imgs, index = _extract_images_from_pdf(path, doc_name, starting_page, force_refresh)
            else:
                imgs = _extract_images_from_docx(path, doc_name, force_refresh)
                for img in imgs:
                    if img.get("rId"):
                        anchor_info = _get_docx_image_anchor(path, img["rId"])
                        img["paragraph_index"] = anchor_info.get("paragraph_index", -1)
                        img["anchor_text"] = anchor_info.get("anchor_text", "")
            all_images.extend(imgs)

    full_text = "\n\n".join(text_parts)
    _log.debug(f"  _read_slices_direct: done, {len(full_text)} chars, {len(found_ids)} slices")

    if extract_images and is_pdf:
        _save_index(doc_name, index)

    if extract_images:
        seen_paths: set[str] = set()
        unique_images = []
        for img in all_images:
            if img["path"] not in seen_paths:
                seen_paths.add(img["path"])
                unique_images.append(img)
        all_images = [img for img in unique_images if img["size"] > 0]
        for img in all_images:
            img["is_small"] = _is_small_image(img["path"])
        all_images = [img for img in all_images if not img["is_small"]]

        ocr_count = 0
        for img in all_images:
            if img["size"] < IMAGE_SIZE_THRESHOLD:
                ocr_text = _ocr_small_image(img["path"])
                if ocr_text:
                    img["ocr"] = ocr_text
                    ocr_count += 1

    written = False
    content_path = None
    if len(full_text) > 0:
        content_path = _get_doc_dir(doc_name) / "output.md"
        md_content = _build_unified_md_output(full_text, all_images, doc_name)
        content_path.write_text(md_content, encoding="utf-8")
        written = True

    progress = []
    if len(found_ids) > 0:
        progress.append(f"已读 {len(found_ids)} 片")
    if all_images:
        progress.append(f"提取图片 {len(all_images)} 张")
    if extract_images and any(img.get("ocr") for img in all_images):
        ocr_n = sum(1 for img in all_images if img.get("ocr"))
        progress.append(f"OCR {ocr_n} 张")

    return {
        "text": full_text,
        "images": all_images,
        "slices": [{"id": sid} for sid in found_ids],
        "paired": False,
        "paired_file": None,
        "written": written,
        "content_path": str(content_path) if written else None,
        "progress": progress,
    }


def _read_single_document(
    file_path: str,
    doc_name: str,
    mode: str = "auto",
    extract_images: bool = True,
    callback_url: str = "",
    force_refresh: bool = False,
) -> dict:
    """
    读取单个文档（PDF 或 DOCX）
    mode: "pdf_only" | "docx_only" | "auto"
    callback_url: 可选，事件通知回调地址
    """
    _log.debug(f"  _read_single: START file={file_path} mode={mode} callback_url={callback_url} mem={mem()}MB")
    _post_callback(callback_url, "document.started", {"doc_name": doc_name, "file_path": file_path})
    _post_callback(callback_url, "read_started", {"file_path": file_path, "doc_name": doc_name, "mode": mode})
    p = Path(file_path)
    ext = p.suffix.lower()

    slices = []
    text_parts = []
    all_images = []
    index = None

    if ext == ".pdf" or (mode in ("auto", "pdf_only") and _is_pdf_by_magic(file_path)):
        slices = _slice_pdf(file_path, doc_name)
        _log.debug(f"  _read_single: {len(slices)} PDF slices, mem={mem()}MB")
        for i, sl in enumerate(slices):
            if i % 5 == 0:
                _log.debug(f"  _read_single: PDF slice {i}/{len(slices)}, mem={mem()}MB")
            txt = _read_pdf_text(sl["path"])
            text_parts.append(f"=== [{sl['id']}] ===\n{txt}")
            if extract_images:
                starting_page = i * SLICE_PAGES + 1
                imgs, index = _extract_images_from_pdf(sl["path"], doc_name, starting_page, force_refresh)
                all_images.extend(imgs)

    elif ext in (".docx", ".doc") or mode == "docx_only":
        slices = _slice_docx(file_path, doc_name)
        _log.debug(f"  _read_single: {len(slices)} DOCX slices, mem={mem()}MB")
        for i, sl in enumerate(slices):
            if i % 5 == 0:
                _log.debug(f"  _read_single: DOCX slice {i}/{len(slices)}, mem={mem()}MB")
            txt = _read_docx_text(sl["path"])
            text_parts.append(f"=== [{sl['id']}] ===\n{txt}")
            if extract_images:
                imgs = _extract_images_from_docx(sl["path"], doc_name, force_refresh)
                for img in imgs:
                    if img.get("rId"):
                        anchor_info = _get_docx_image_anchor(sl["path"], img["rId"])
                        img["paragraph_index"] = anchor_info.get("paragraph_index", -1)
                        img["anchor_text"] = anchor_info.get("anchor_text", "")
                all_images.extend(imgs)

    else:
        txt = _read_generic_text(file_path)
        text_parts.append(txt)

    if extract_images and index is not None:
        _save_index(doc_name, index)

    full_text = "\n\n".join(text_parts)
    _log.debug(f"  _read_single: text done, {len(full_text)} chars, mem={mem()}MB")

    ocr_count = 0
    if extract_images:
        seen_paths: set[str] = set()
        unique_images = []
        for img in all_images:
            if img["path"] not in seen_paths:
                seen_paths.add(img["path"])
                unique_images.append(img)
        all_images = [img for img in unique_images if img["size"] > 0]
        for img in all_images:
            img["is_small"] = _is_small_image(img["path"])
        all_images = [img for img in all_images if not img["is_small"]]
        _log.debug(f"  _read_single: OCR on {len(all_images)} unique images")
        for idx, img in enumerate(all_images):
            if img["size"] < IMAGE_SIZE_THRESHOLD:
                ocr_text = _ocr_small_image(img["path"])
                if ocr_text:
                    img["ocr"] = ocr_text
                    ocr_count += 1
        _log.debug(f"  _read_single: OCR done, {ocr_count} succeeded, mem={mem()}MB")

    written = False
    content_path = None
    if len(full_text) > 0:
        content_path = _get_doc_dir(doc_name) / "output.md"
        md_content = _build_unified_md_output(full_text, all_images, doc_name, file_path)
        content_path.write_text(md_content, encoding="utf-8")
        written = True

    progress = []
    progress.append(f"分片 {len(slices)} 片")
    if all_images:
        progress.append(f"提取图片 {len(all_images)} 张")
    if extract_images and any(img.get("ocr") for img in all_images):
        ocr_n = sum(1 for img in all_images if img.get("ocr"))
        progress.append(f"OCR {ocr_n} 张")

    _log.debug(f"  _read_single: DONE mem={mem()}MB")
    _post_callback(callback_url, "document.completed", {
        "doc_name": doc_name,
        "chars": len(full_text),
        "slices": len(slices),
        "images": len(all_images),
        "ocr_count": ocr_count,
    })
    return {
        "text": full_text,
        "images": all_images,
        "slices": slices,
        "paired": False,
        "paired_file": None,
        "written": written,
        "content_path": str(content_path) if written else None,
        "progress": progress,
    }


def _read_paired_documents(
    pdf_path: str, docx_path: str, extract_images: bool = True, callback_url: str = "", force_refresh: bool = False
) -> dict:
    doc_name = _make_doc_name(pdf_path)
    _log.debug(f"  _read_paired: START pdf={doc_name} callback_url={callback_url} mem={mem()}MB")
    _post_callback(callback_url, "document.started", {"doc_name": doc_name, "pdf_path": pdf_path, "docx_path": docx_path})
    _post_callback(callback_url, "read_started", {"pdf_path": pdf_path, "docx_path": docx_path, "doc_name": doc_name})

    _log.debug(f"  _read_paired: slicing PDF")
    pdf_slices = _slice_pdf(pdf_path, doc_name)
    print(f"  _read_paired: PDF sliced into {len(pdf_slices)} slices, mem={resource.getrusage(resource.RUSAGE_SELF).ru_maxrss // 1024}MB", flush=True)
    _log.debug(f"  _read_paired: {len(pdf_slices)} PDF slices, mem={mem()}MB")

    all_images = []
    if extract_images:
        _log.debug(f"  _read_paired: extracting images from {len(pdf_slices)} slices")
        for i, sl in enumerate(pdf_slices):
            if i % 5 == 0:
                _log.debug(f"  _read_paired: image slice {i}/{len(pdf_slices)}, mem={mem()}MB")
            starting_page = i * SLICE_PAGES + 1
            imgs, index = _extract_images_from_pdf(sl["path"], doc_name, starting_page, force_refresh)
            all_images.extend(imgs)
            print(f"  _read_paired: slice {i+1}/{len(pdf_slices)} done: {len(imgs)} imgs, total={len(all_images)}, mem={resource.getrusage(resource.RUSAGE_SELF).ru_maxrss // 1024}MB", flush=True)

    if extract_images:
        _save_index(doc_name, index)

    _log.debug(f"  _read_paired: slicing DOCX")
    docx_slices = _slice_docx(docx_path, doc_name)
    print(f"  _read_paired: DOCX sliced into {len(docx_slices)} slices, mem={resource.getrusage(resource.RUSAGE_SELF).ru_maxrss // 1024}MB", flush=True)
    _log.debug(f"  _read_paired: {len(docx_slices)} DOCX slices, mem={mem()}MB")

    text_parts = []
    for i, sl in enumerate(docx_slices):
        if i % 5 == 0:
            _log.debug(f"  _read_paired: DOCX slice {i}/{len(docx_slices)}, mem={mem()}MB")
        txt = _read_docx_text(sl["path"])
        text_parts.append(f"=== [{sl['id']}] ===\n{txt}")

    full_text = "\n\n".join(text_parts)
    _log.debug(f"  _read_paired: text done, {len(full_text)} chars, mem={mem()}MB")

    ocr_count = 0
    if extract_images:
        seen_paths: set[str] = set()
        unique_images = []
        for img in all_images:
            if img["path"] not in seen_paths:
                seen_paths.add(img["path"])
                unique_images.append(img)
        all_images = [img for img in unique_images if img["size"] > 0]
        for img in all_images:
            img["is_small"] = _is_small_image(img["path"])
        all_images = [img for img in all_images if not img["is_small"]]
        print(f"  _read_paired: after dedup: {len(all_images)} unique images, mem={resource.getrusage(resource.RUSAGE_SELF).ru_maxrss // 1024}MB", flush=True)
        _log.debug(f"  _read_paired: OCR on {len(all_images)} unique images")
        for img in all_images:
            if img["size"] < IMAGE_SIZE_THRESHOLD:
                ocr_text = _ocr_small_image(img["path"])
                if ocr_text:
                    img["ocr"] = ocr_text
                    ocr_count += 1
        _log.debug(f"  _read_paired: OCR done, {ocr_count} succeeded, mem={mem()}MB")

    written = False
    content_path = None
    if len(full_text) > 0:
        content_path = _get_doc_dir(doc_name) / "output.md"
        print(f"  _read_paired: about to build MD: {len(all_images)} images, {len(full_text)} text chars, mem={resource.getrusage(resource.RUSAGE_SELF).ru_maxrss // 1024}MB", flush=True)
        md_content = _build_unified_md_output(full_text, all_images, doc_name, pdf_path)
        content_path.write_text(md_content, encoding="utf-8")
        written = True

    progress = []
    progress.append(f"PDF {len(pdf_slices)} 片")
    progress.append(f"DOCX {len(docx_slices)} 片")
    if all_images:
        progress.append(f"提取图片 {len(all_images)} 张")
    if extract_images and any(img.get("ocr") for img in all_images):
        ocr_n = sum(1 for img in all_images if img.get("ocr"))
        progress.append(f"OCR {ocr_n} 张")

    _log.debug(f"  _read_paired: DONE mem={mem()}MB")
    _post_callback(callback_url, "document.completed", {
        "doc_name": doc_name,
        "chars": len(full_text),
        "pdf_slices": len(pdf_slices),
        "docx_slices": len(docx_slices),
        "images": len(all_images),
        "ocr_count": ocr_count,
    })
    return {
        "text": full_text,
        "images": all_images,
        "pdf_slices": pdf_slices,
        "docx_slices": docx_slices,
        "slices": docx_slices,
        "paired": True,
        "paired_file": docx_path,
        "written": written,
        "content_path": str(content_path) if written else None,
        "progress": progress,
    }


def _build_unified_md_output(text: str, images: list[dict], doc_name: str, doc_path: str = None) -> str:
    """
    构建统一MD输出，包含图片锚定信息。
    图片按位置插入到文本中，而非仅在末尾追加。
    格式：
      ![](path/to/img.png){.positioned page=1 y=245}
      Image: OCR result or "see nearby content"
    """
    import re
    # 按 page 和 y 坐标排序图片，以便按序插入
    sorted_images = sorted(images, key=lambda img: (img.get("page", 0), img.get("y", 0)))

    # 将文本按页分割（=== [px-y] === 格式标记）
    page_sections = []
    current_section = {"page": None, "lines": []}
    for line in text.split("\n"):
        # 检测页标记：=== [px-y] ===
        m = re.match(r"=== \[p(\d+)-\d+\] ===", line)
        if m:
            if current_section["lines"]:
                page_sections.append(current_section)
            current_section = {"page": int(m.group(1)), "lines": [line]}
        else:
            current_section["lines"].append(line)
    if current_section["lines"]:
        page_sections.append(current_section)

    # 构建增量MD：文本 + 图片交替
    result_lines = []
    img_idx = 0

    for section in page_sections:
        section_page = section["page"] if section["page"] is not None else 0
        section_text = "\n".join(section["lines"])
        result_lines.append(section_text)

        # 插入属于该页的图片
        while img_idx < len(sorted_images):
            img = sorted_images[img_idx]
            img_page = img.get("page") or 0
            if img_page > section_page:
                break  # 等待下一页

            if img_page is not None and img_page == section_page:
                img_path = img.get("path", "")
                img_y = img.get("y", 0)
                ocr_text = img.get("ocr", "")
                img_anchor = img.get("anchor_text", "")
                is_small = img.get("is_small", False)
                pos_str = f"page={img_page}" if img_page is not None else "page=0"
                if img_y:
                    pos_str += f" y={int(img_y)}"
                anchor_text = ocr_text if ocr_text else img_anchor
                # Get nearby text as anchor if no OCR or anchor
                if not anchor_text and section_text:
                    nearby = section_text.strip()[:50].replace('\n', ' ')
                    anchor_text = f"Image ({nearby}...)" if nearby else "Image"
                elif not anchor_text:
                    anchor_text = f"Image on page {img_page}" if img_page else "Image"
                # Small image override
                if is_small and not ocr_text and not img_anchor:
                    anchor_text = f"Small image on page {img_page}" if img_page else "Small image"
                result_lines.append(f"![]({img_path}){{.positioned {pos_str}}}")
                result_lines.append(f"Image: {anchor_text}\n")
            img_idx += 1

    # 追加剩余未插入的图片（无页信息或页号超出）
    while img_idx < len(sorted_images):
        img = sorted_images[img_idx]
        img_path = img.get("path", "")
        page = img.get("page", 0)
        img_y = img.get("y", 0)
        ocr_text = img.get("ocr", "")
        img_anchor = img.get("anchor_text", "")
        is_small = img.get("is_small", False)
        pos_str = f"page={page}" if page else "page=0"
        if img_y:
            pos_str += f" y={int(img_y)}"
        anchor_text = ocr_text if ocr_text else img_anchor
        # Get nearby text as anchor if no OCR or anchor
        if not anchor_text and section_text:
            nearby = section_text.strip()[:50].replace('\n', ' ')
            anchor_text = f"Image ({nearby}...)" if nearby else "Image"
        elif not anchor_text:
            anchor_text = f"Image on page {page}" if page else "Image"
        # Small image override
        if is_small and not ocr_text and not img_anchor:
            anchor_text = f"Small image on page {page}" if page else "Small image"
        result_lines.append(f"![]({img_path}){{.positioned {pos_str}}}")
        result_lines.append(f"Image: {anchor_text}\n")
        img_idx += 1

    return "\n".join(result_lines)


def _update_image_ocr(doc_name: str, img_id: str, ocr_result: str, position_info: dict) -> bool:
    """更新单张图片的OCR结果到index.json，返回是否成功"""
    lock_path = _get_doc_dir(doc_name) / ".index.lock"
    lock_fd = _lock_file(lock_path)
    try:
        index = _load_index(doc_name)
        img_found = False
        for img in index.get("images", []):
            if img.get("name") == img_id or img.get("md5") == img_id:
                img["ocr_status"] = "completed"
                img["ocr_result"] = ocr_result
                img["ocr_source"] = "manual"
                if position_info:
                    img["position_info"] = position_info
                img_found = True
                break
        if not img_found:
            return False
        _save_index_nolock(doc_name, index)
        return True
    finally:
        _unlock_file(lock_fd)


def _rebuild_output_md(doc_name: str):
    """根据index中的OCR结果重建output.md"""
    index = _load_index(doc_name)
    doc_dir = _get_doc_dir(doc_name)
    content_path = doc_dir / "output.md"

    if not content_path.exists():
        return

    existing_text = content_path.read_text(encoding="utf-8")

    ocr_map = {}
    for img in index.get("images", []):
        if img.get("ocr_result"):
            img_path = img.get("path", "")
            ocr_map[img_path] = img.get("ocr_result", "")

    lines = existing_text.split("\n")
    result_lines = []
    for line in lines:
        if line.strip().startswith("Image:") and not line.strip().endswith("see nearby content"):
            result_lines.append(line)
        else:
            result_lines.append(line)

    content_path.write_text("\n".join(result_lines), encoding="utf-8")
