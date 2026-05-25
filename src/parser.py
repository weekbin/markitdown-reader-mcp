# -*- coding: utf-8 -*-
"""PDF/DOCX parsing functions extracted from server.py"""

from __future__ import annotations

import hashlib
import json
import zipfile
import copy
import time
import re
import logging
from pathlib import Path
from typing import Optional

from src.constants import SLICE_PAGES, SLICE_BLOCKS, BASE_DIR

_log = logging.getLogger("markitdown")


def mem():
    import resource
    return resource.getrusage(resource.RUSAGE_SELF).ru_maxrss // 1024


# ─────────────────────────────────────────────────────────────────
# Storage helpers (forwarded from server.py until src/storage exists)
# ─────────────────────────────────────────────────────────────────

def _get_doc_dir(doc_name: str) -> Path:
    """获取文档专属目录，不存在则创建"""
    d = BASE_DIR / doc_name
    d.mkdir(parents=True, exist_ok=True)
    return d


def _get_slices_dir(doc_name: str) -> Path:
    d = _get_doc_dir(doc_name) / "slices"
    d.mkdir(parents=True, exist_ok=True)
    return d


def _get_images_dir(doc_name: str) -> Path:
    d = _get_doc_dir(doc_name) / "images"
    d.mkdir(parents=True, exist_ok=True)
    return d


def _get_index_path(doc_name: str) -> Path:
    return _get_doc_dir(doc_name) / "index.json"


def _load_index(doc_name: str) -> dict:
    p = _get_index_path(doc_name)
    if p.exists():
        data = json.loads(p.read_text(encoding="utf-8"))
        data.setdefault("content_hash", "")
        data.setdefault("source_file", "")
        data.setdefault("runs", [])
        return data
    return {"slices": [], "images": [], "ocr_done": [], "content_hash": "", "source_file": "", "runs": []}


def _save_index(doc_name: str, index: dict):
    import fcntl

    def _lock_file(lock_path: Path):
        lock_path.parent.mkdir(parents=True, exist_ok=True)
        lock_fd = open(lock_path, "w")
        fcntl.flock(lock_fd.fileno(), fcntl.LOCK_EX)
        return lock_fd

    def _unlock_file(lock_fd):
        fcntl.flock(lock_fd.fileno(), fcntl.LOCK_UN)
        lock_fd.close()

    lock_path = _get_doc_dir(doc_name) / ".index.lock"
    lock_fd = _lock_file(lock_path)
    try:
        p = _get_index_path(doc_name)
        p.write_text(json.dumps(index, ensure_ascii=False, indent=2), encoding="utf-8")
    finally:
        _unlock_file(lock_fd)


# ─────────────────────────────────────────────────────────────────
# PDF 分片
# ─────────────────────────────────────────────────────────────────

def _slice_pdf(src: str, doc_name: str, pages_per_slice: int = SLICE_PAGES, time_budget: float = 300.0) -> list[dict]:
    """PyMuPDF 按页分片，返回分片信息列表。支持时间预算控制，增量保存进度。"""
    import fitz
    _log.debug(f"  _slice_pdf: opening {src}")
    start_time = time.monotonic()
    with fitz.open(src) as doc:
        total = len(doc)
        _log.debug(f"  _slice_pdf: {total} pages, mem={mem()}MB")
        slices_dir = _get_slices_dir(doc_name)
        slices = []
        for i in range(0, len(doc), pages_per_slice):
            elapsed = time.monotonic() - start_time
            if elapsed >= time_budget:
                _log.debug(f"  _slice_pdf: time budget exceeded at slice {i // pages_per_slice}, saving partial progress")
                break
            end = min(i + pages_per_slice, len(doc))
            slice_doc = fitz.open()
            try:
                slice_doc.insert_pdf(doc, from_page=i, to_page=end - 1)
                out_path = slices_dir / f"slice_{i // pages_per_slice:03d}.pdf"
                slice_doc.save(str(out_path))
            finally:
                slice_doc.close()
            slices.append({
                "id": f"p{i + 1}-{end}",
                "path": str(out_path),
                "pages": (i + 1, end),
                "total": total
            })
            if (i // pages_per_slice) % 5 == 0:
                index = _load_index(doc_name)
                existing_docx = index.get("docx_slices", index.get("slices", []))
                index["pdf_slices"] = slices
                index["docx_slices"] = existing_docx
                _save_index(doc_name, index)
    index = _load_index(doc_name)
    existing_docx = index.get("docx_slices", index.get("slices", []))
    index["pdf_slices"] = slices
    index["docx_slices"] = existing_docx
    _save_index(doc_name, index)
    _log.debug(f"  _slice_pdf: done, {len(slices)} slices, index saved, elapsed={time.monotonic() - start_time:.1f}s")
    return slices


# ─────────────────────────────────────────────────────────────────
# DOCX 分片
# ─────────────────────────────────────────────────────────────────

def _slice_docx(src: str, doc_name: str, blocks_per_slice: int = SLICE_BLOCKS, time_budget: float = 300.0) -> list[dict]:
    """python-docx XML 按 block 分片，返回分片信息列表。支持时间预算控制，增量保存进度。"""
    from lxml import etree
    _log.debug(f"  _slice_docx: opening {src}")
    start_time = time.monotonic()
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    slices_dir = _get_slices_dir(doc_name)

    with zipfile.ZipFile(src, "r") as z:
        doc_xml = z.read("word/document.xml")
        # 只读取非媒体文件的元数据，媒体文件延迟加载
        all_items = {}
        for item in z.namelist():
            if item.startswith("word/media/") or item.endswith(".png") or item.endswith(".jpg") or item.endswith(".jpeg") or item.endswith(".gif") or item.endswith(".bmp"):
                continue  # 延迟加载大媒体文件
            all_items[item] = z.read(item)
    tree = etree.fromstring(doc_xml)
    body = tree.find(f"{{{W}}}body")

    blocks = []
    for child in body:
        tag = child.tag.split("}")[-1]
        if tag in ("p", "tbl"):
            blocks.append(child)
    _log.debug(f"  _slice_docx: {len(blocks)} blocks, mem={mem()}MB")

    slices = []
    total_slices = (len(blocks) + blocks_per_slice - 1) // blocks_per_slice
    for i in range(0, len(blocks), blocks_per_slice):
        elapsed = time.monotonic() - start_time
        if elapsed >= time_budget:
            _log.debug(f"  _slice_docx: time budget exceeded at slice {i // blocks_per_slice}, saving partial progress")
            break
        slice_idx = i // blocks_per_slice
        if slice_idx % 10 == 0:
            _log.debug(f"  _slice_docx: slice {slice_idx}/{total_slices}, mem={mem()}MB")
        end = min(i + blocks_per_slice, len(blocks))
        out_path = slices_dir / f"slice_{slice_idx:03d}.docx"

        tree = etree.fromstring(doc_xml)
        body = tree.find(f"{{{W}}}body")
        for child in list(body):
            body.remove(child)
        for block in blocks[i:end]:
            body.append(copy.deepcopy(block))
        with zipfile.ZipFile(str(out_path), "w", zipfile.ZIP_DEFLATED) as zout:
            for item, data in all_items.items():
                if item == "word/document.xml":
                    zout.writestr(item, etree.tostring(tree))
                else:
                    zout.writestr(item, data)
            # 媒体文件延迟写入
            with zipfile.ZipFile(src, "r") as zsrc:
                for item in zsrc.namelist():
                    if item.startswith("word/media/"):
                        zout.writestr(item, zsrc.read(item))

        slices.append({
            "id": f"b{i + 1}-{end}",
            "path": str(out_path),
            "blocks": (i + 1, end),
            "total": len(blocks)
        })
    index = _load_index(doc_name)
    existing_pdf = index.get("pdf_slices", [])
    index["pdf_slices"] = existing_pdf
    index["docx_slices"] = slices
    _save_index(doc_name, index)
    _log.debug(f"  _slice_docx: done, {len(slices)} slices, index saved, elapsed={time.monotonic() - start_time:.1f}s, mem={mem()}MB")
    return slices


# ─────────────────────────────────────────────────────────────────
# 文字提取
# ─────────────────────────────────────────────────────────────────

def _read_pdf_text(path: str) -> str:
    """PyMuPDF 提取 PDF 文字"""
    import fitz
    _log.debug(f"  _read_pdf_text: opening {path}")
    with fitz.open(path) as doc:
        total = len(doc)
        parts = []
        for i, page in enumerate(doc):
            if i % 20 == 0:
                _log.debug(f"  _read_pdf_text: page {i+1}/{total}")
            text = page.get_text("text")
            if text.strip():
                parts.append(f"--- 第 {i + 1} 页 ---\n{text}")
        _log.debug(f"  _read_pdf_text: done, {len(parts)} pages with text")
        return "\n\n".join(parts)


def _get_image_position_info(pdf_path: str, page_num: int, bbox: tuple) -> dict:
    """获取图片在PDF中的位置信息，返回最近文字"""
    import fitz
    try:
        with fitz.open(pdf_path) as doc:
            if page_num < 1 or page_num > len(doc):
                return {"page": page_num, "y": 0, "nearest_text_above": "", "nearest_text_below": ""}
            page = doc[page_num - 1]
            text_dict = page.get_text("dict")
            blocks = text_dict.get("blocks", [])

            img_y = bbox[1] if bbox else 0
            nearest_above = ""
            nearest_below = ""
            min_dist_above = float("inf")
            min_dist_below = float("inf")

            for block in blocks:
                if block.get("type") == 0:
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            span_y = span.get("origin", (0, 0))[1]
                            span_text = span.get("text", "").strip()
                            if not span_text:
                                continue
                            dist = img_y - span_y
                            if dist > 0 and dist < min_dist_above:
                                min_dist_above = dist
                                nearest_above = span_text
                            elif dist < 0 and -dist < min_dist_below:
                                min_dist_below = -dist
                                nearest_below = span_text

            return {
                "page": page_num,
                "y": img_y,
                "nearest_text_above": nearest_above,
                "nearest_text_below": nearest_below,
            }
    except Exception:
        return {"page": page_num, "y": 0, "nearest_text_above": "", "nearest_text_below": ""}


def _read_docx_text(path: str) -> str:
    """python-docx 提取 DOCX 文字+表格"""
    from docx import Document
    _log.debug(f"  _read_docx_text: opening {path}")
    doc = Document(path)
    parts = []
    total = len(list(doc.element.body))
    _log.debug(f"  _read_docx_text: {total} blocks")

    for idx, elem in enumerate(doc.element.body):
        if idx % 100 == 0:
            _log.debug(f"  _read_docx_text: block {idx}/{total}")
        tag = elem.tag.split("}")[-1]
        if tag == "p":
            text = "".join(
                t.text or "" for t in elem.iter(f"{{{docx_ns}}}t")
            )
            if text.strip():
                parts.append(text)
        elif tag == "tbl":
            tbl_text = _docx_table_to_markdown(elem)
            if tbl_text:
                parts.append(tbl_text)

    _log.debug(f"  _read_docx_text: done, {len(parts)} blocks with content")
    return "\n\n".join(parts)


def _docx_table_to_markdown(tbl_elem) -> str:
    """将 DOCX 表格元素转为 Markdown 表格"""
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    rows = tbl_elem.findall(f".//{{{W}}}tr")
    if not rows:
        return ""

    parts = []
    for ri, tr in enumerate(rows):
        cells = tr.findall(f".//{{{W}}}tc")
        row_text = []
        for tc in cells:
            cell_text = "".join(
                t.text or "" for t in tc.iter(f"{{{W}}}t")
            ).strip()
            row_text.append(cell_text)
        parts.append("| " + " | ".join(row_text) + " |")
        if ri == 0:
            parts.append("| " + " | ".join(["---"] * len(row_text)) + " |")
    return "\n".join(parts)


docx_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ─────────────────────────────────────────────────────────────────
# 图片提取（统一入口）
# ─────────────────────────────────────────────────────────────────

def _resize_image_if_needed(img_bytes: bytes, max_dim: int = 1280) -> bytes:
    """
    如果图片宽或高超过 max_dim 像素，缩图后返回新 bytes。
    否则直接返回原始 bytes。
    使用 PIL LANCZOS 重采样，保持灰度/RGB 模式。
    """
    try:
        from PIL import Image
        import io
        img = Image.open(io.BytesIO(img_bytes))
        w, h = img.size
        if w <= max_dim and h <= max_dim:
            return img_bytes  # 不需要缩图
        # 等比缩放
        img.thumbnail((max_dim, max_dim), Image.LANCZOS)
        out = io.BytesIO()
        # 保持原始格式；GIF 需要特殊处理保持动画
        save_fmt = img.format if img.format in ("PNG", "JPEG", "GIF", "BMP", "WEBP") else "PNG"
        if img.mode == "RGBA" and save_fmt == "JPEG":
            img = img.convert("RGB")
        img.save(out, format=save_fmt, optimize=True)
        return out.getvalue()
    except Exception:
        return img_bytes  # 失败时保原图


def _extract_images_from_pdf(pdf_path: str, doc_name: str, starting_page: int = 1, force_refresh: bool = False) -> tuple[list[dict], dict]:
    """从 PDF 提取图片，返回图片信息列表（含大小/页码/MD5/bbox/local_page）"""
    import fitz
    _log.debug(f"  _extract_images_from_pdf: opening {pdf_path}")
    images_dir = _get_images_dir(doc_name)
    index = _load_index(doc_name)
    if force_refresh:
        seen_hashes: set[str] = set()
        index["images"] = []
    else:
        seen_hashes = set(img["md5"] for img in index["images"])
    results = []

    with fitz.open(pdf_path) as doc:
        total_pages = len(doc)
        _log.debug(f"  _extract_images_from_pdf: {total_pages} pages, mem={mem()}MB")
        for page_num in range(len(doc)):
            if page_num % 20 == 0:
                _log.debug(f"  _extract_images_from_pdf: page {page_num+1}/{total_pages}, mem={mem()}MB")
            page = doc[page_num]
            img_list = page.get_images(full=True)
            img_rects = {}
            # Get image positions using xref as key
            for img in img_list:
                xref = img[0]
                rects = page.get_image_rects(xref)
                if rects:
                    img_rects[xref] = rects[0].rect if hasattr(rects[0], 'rect') else rects[0]

            # Get text blocks for position info ONCE per page (avoid re-opening)
            text_dict = page.get_text("dict")
            blocks = text_dict.get("blocks", [])

            original_page = starting_page + page_num
            for img_index, img in enumerate(img_list):
                xref = img[0]
                base_image = doc.extract_image(xref)
                img_bytes = base_image["image"]
                content_hash = hashlib.md5(img_bytes).hexdigest()
                ext = base_image.get("ext", "png").lower()
                if ext not in ("png", "jpg", "jpeg", "gif", "bmp", "webp"):
                    ext = "png"
                img_hash = content_hash[:12]
                out_name = f"{doc_name}_p{original_page}_i{img_index}_{img_hash}.{ext}"
                out_path = images_dir / out_name

                bbox = img_rects.get(xref)
                bbox_tuple = (bbox.x0, bbox.y0, bbox.x1, bbox.y1) if bbox else None

                # Compute position info from cached text blocks (avoid re-opening PDF)
                img_y = bbox.y0 if bbox else 0
                nearest_above = ""
                nearest_below = ""
                min_dist_above = float("inf")
                min_dist_below = float("inf")
                for block in blocks:
                    if block.get("type") == 0:
                        for line in block.get("lines", []):
                            for span in line.get("spans", []):
                                span_y = span.get("origin", (0, 0))[1]
                                span_text = span.get("text", "").strip()
                                if not span_text:
                                    continue
                                dist = img_y - span_y
                                if dist > 0 and dist < min_dist_above:
                                    min_dist_above = dist
                                    nearest_above = span_text
                                elif dist < 0 and -dist < min_dist_below:
                                    min_dist_below = -dist
                                    nearest_below = span_text

                is_new = content_hash not in seen_hashes
                if is_new:
                    seen_hashes.add(content_hash)
                    img_bytes = _resize_image_if_needed(img_bytes)
                    out_path.write_bytes(img_bytes)
                    index["images"].append({
                        "name": out_name,
                        "md5": content_hash,
                        "path": str(out_path),
                        "page": original_page,
                        "size": len(img_bytes)
                    })

                results.append({
                    "name": out_name,
                    "path": str(out_path),
                    "page": original_page,
                    "size": len(img_bytes),
                    "is_new": is_new,
                    "xref": xref,
                    "bbox": bbox_tuple,
                    "local_page": page_num + 1,
                    "y": img_y,
                    "nearest_text_above": nearest_above,
                    "nearest_text_below": nearest_below,
                })

    _log.debug(f"  _extract_images_from_pdf: done, {len(results)} images, mem={mem()}MB")
    return results, index


def _extract_images_from_docx(docx_path: str, doc_name: str, force_refresh: bool = False) -> list[dict]:
    """从 DOCX 提取图片，返回图片信息列表（含rId）"""
    _log.debug(f"  _extract_images_from_docx: opening {docx_path}")
    images_dir = _get_images_dir(doc_name)
    index = _load_index(doc_name)
    if force_refresh:
        seen_hashes: set[str] = set()
        index["images"] = []
    else:
        seen_hashes = set(img["md5"] for img in index["images"])
    results = []

    with zipfile.ZipFile(docx_path, "r") as z:
        media_files = [f for f in z.namelist() if f.startswith("word/media/")]
        rels_files = [f for f in z.namelist() if f.endswith(".rels")]
        rid_to_image = {}
        for rels_file in rels_files:
            rels_content = z.read(rels_file).decode("utf-8")
            for match in re.finditer(r'Id="(rId\d+)"[^>]*Target="media/([^"]+)"', rels_content):
                rid = match.group(1)
                image_path = "word/" + match.group(2)
                rid_to_image[image_path] = rid
        _log.debug(f"  _extract_images_from_docx: {len(media_files)} media files, mem={mem()}MB")
        for i, mf in enumerate(media_files):
            if i % 50 == 0:
                _log.debug(f"  _extract_images_from_docx: {i}/{len(media_files)}")
            img_bytes = z.read(mf)
            content_hash = hashlib.md5(img_bytes).hexdigest()
            ext = Path(mf).suffix.lower()
            if ext not in (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"):
                ext = ".png"
            img_hash = content_hash[:12]
            out_name = f"{doc_name}_docx_{img_hash}{ext}"
            out_path = images_dir / out_name
            rId = rid_to_image.get(mf, "")

            is_new = content_hash not in seen_hashes
            if is_new:
                seen_hashes.add(content_hash)
                img_bytes = _resize_image_if_needed(img_bytes)
                out_path.write_bytes(img_bytes)
                index["images"].append({
                    "name": out_name,
                    "md5": content_hash,
                    "path": str(out_path),
                    "size": len(img_bytes)
                })

            results.append({
                "name": out_name,
                "path": str(out_path),
                "size": len(img_bytes),
                "is_new": is_new,
                "rId": rId,
            })

    _log.debug(f"  _extract_images_from_docx: done, {len(results)} images, mem={mem()}MB")
    return results


def _get_docx_image_anchor(docx_path: str, image_rId: str) -> dict:
    """获取DOCX中图片关联的段落位置信息"""
    try:
        from docx import Document
        doc = Document(docx_path)
        for idx, para in enumerate(doc.paragraphs):
            for run in para.runs:
                if hasattr(run, "_element"):
                    inline_shapes = run._element.findall(
                        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}inline"
                    )
                    for shape in inline_shapes:
                        docPr = shape.find(
                            ".//{http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing}blip"
                        )
                        if docPr is not None:
                            rId = docPr.get(
                                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                            )
                            if rId == image_rId:
                                return {
                                    "paragraph_index": idx,
                                    "anchor_text": para.text.strip()[:100],
                                }
        for idx, para in enumerate(doc.paragraphs):
            if image_rId in para._element.xml:
                return {
                    "paragraph_index": idx,
                    "anchor_text": para.text.strip()[:100],
                }
    except Exception:
        pass
    return {"paragraph_index": -1, "anchor_text": ""}


# ─────────────────────────────────────────────────────────────────
# 辅助
# ─────────────────────────────────────────────────────────────────

def _is_pdf_by_magic(path: str) -> bool:
    """通过文件魔数判断是否为 PDF"""
    try:
        with open(path, "rb") as f:
            header = f.read(8)
        return header.startswith(b"%PDF")
    except Exception:
        return False


def _read_generic_text(path: str) -> str:
    """通用文本读取：先尝试 PDF，若失败则尝试 DOCX"""
    if _is_pdf_by_magic(path):
        return _read_pdf_text(path)
    elif path.lower().endswith(".docx"):
        return _read_docx_text(path)
    else:
        _log.debug(f"  _read_generic_text: unsupported file type: {path}")
        return ""
