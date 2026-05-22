# -*- coding: utf-8 -*-
# @Author:  (AI Assistant)
# @Time:    2026-05-21
"""
markitdown-reader MCP Service
- PDF: PyMuPDF 提取文字 + 图片
- DOCX: markitdown 提取文字 + 表格
- 支持配对文件检测（PDF+DOXC 同时处理）
- 自动分片（5页/片）
- 持久化图片池（~/.opencode/markitdown/）
- 小图内置 OCR
"""

from __future__ import annotations

import os
import re
import json
import hashlib
import zipfile
import tempfile
import logging
import traceback
from pathlib import Path
from typing import Optional

LOG_FILE = Path.home() / ".opencode" / "markitdown" / f"server_{os.getpid()}.log"
LOG_FILE.parent.mkdir(parents=True, exist_ok=True)

_log = logging.getLogger("markitdown")
_log.setLevel(logging.DEBUG)
_handler = logging.FileHandler(LOG_FILE, encoding="utf-8")
_handler.setFormatter(logging.Formatter("%(asctime)s %(message)s", datefmt="%H:%M:%S"))
_log.addHandler(_handler)

def mem():
    import resource
    return resource.getrusage(resource.RUSAGE_SELF).ru_maxrss // 1024

from mcp.server.fastmcp import FastMCP

mcp = FastMCP("markitdown-reader")

BASE_DIR = Path.home() / ".opencode" / "markitdown"
SLICE_PAGES = 5  # 每5页一片
SLICE_BLOCKS = 200  # DOCX 每200 block 一片
IMAGE_SIZE_THRESHOLD = 50 * 1024  # 50KB
MAX_CHARS_RETURN = 400_000  # 直接返回的最大字符数


# ─────────────────────────────────────────────────────────────────
# 辅助：持久化目录
# ─────────────────────────────────────────────────────────────────

def _get_doc_dir(doc_name: str) -> Path:
    """获取文档专属目录，不存在则创建"""
    d = BASE_DIR / doc_name
    d.mkdir(parents=True, exist_ok=True)
    return d


def _get_slices_dir(doc_name: str) -> Path:
    d = _get_doc_dir(doc_name) / "slices"
    d.mkdir(parents=True, exist_ok=True)
    return d


def _get_images_dir(doc_name: str) -> Path:
    d = _get_doc_dir(doc_name) / "images"
    d.mkdir(parents=True, exist_ok=True)
    return d


def _get_index_path(doc_name: str) -> Path:
    return _get_doc_dir(doc_name) / "index.json"


def _load_index(doc_name: str) -> dict:
    p = _get_index_path(doc_name)
    if p.exists():
        return json.loads(p.read_text(encoding="utf-8"))
    return {"slices": [], "images": [], "ocr_done": []}


def _save_index(doc_name: str, index: dict):
    p = _get_index_path(doc_name)
    p.write_text(json.dumps(index, ensure_ascii=False, indent=2), encoding="utf-8")


# ─────────────────────────────────────────────────────────────────
# 配对文件检测
# ─────────────────────────────────────────────────────────────────

def _find_paired_file(path: str) -> Optional[str]:
    """
    在同目录查找同名不同扩展名的配对文件。

    跳过已经是分片的文件（path 包含 /slices/ 或文件名含 slice_ 前缀），
    避免对分片文件错误触发配对逻辑导致重复分片和无限递归。
    """
    p = Path(path)
    if "/slices/" in str(p) or p.name.startswith("slice_"):
        return None
    stem = p.stem
    directory = p.parent
    for ext in [".pdf", ".docx"]:
        if p.suffix.lower() != ext:
            paired = directory / f"{stem}{ext}"
            if paired.exists():
                return str(paired)
    return None


# ─────────────────────────────────────────────────────────────────
# PDF 分片
# ─────────────────────────────────────────────────────────────────

def _slice_pdf(src: str, doc_name: str, pages_per_slice: int = SLICE_PAGES) -> list[dict]:
    """PyMuPDF 按页分片，返回分片信息列表"""
    import fitz
    _log.debug(f"  _slice_pdf: opening {src}")
    with fitz.open(src) as doc:
        total = len(doc)
        _log.debug(f"  _slice_pdf: {total} pages, mem={mem()}MB")
        slices_dir = _get_slices_dir(doc_name)
        slices = []
        for i in range(0, len(doc), pages_per_slice):
            end = min(i + pages_per_slice, len(doc))
            slice_doc = fitz.open()
            try:
                slice_doc.insert_pdf(doc, from_page=i, to_page=end - 1)
                out_path = slices_dir / f"slice_{i // pages_per_slice:03d}.pdf"
                slice_doc.save(str(out_path))
                slices.append({
                    "id": f"p{i + 1}-{end}",
                    "path": str(out_path),
                    "pages": (i + 1, end),
                    "total": total
                })
            finally:
                slice_doc.close()
        _log.debug(f"  _slice_pdf: done, {len(slices)} slices")
        return slices


# ─────────────────────────────────────────────────────────────────
# DOCX 分片
# ─────────────────────────────────────────────────────────────────

def _slice_docx(src: str, doc_name: str, blocks_per_slice: int = SLICE_BLOCKS) -> list[dict]:
    """python-docx XML 按 block 分片，返回分片信息列表"""
    import zipfile, copy
    from lxml import etree
    _log.debug(f"  _slice_docx: opening {src}")
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    slices_dir = _get_slices_dir(doc_name)

    with zipfile.ZipFile(src, "r") as z:
        doc_xml = z.read("word/document.xml")
    tree = etree.fromstring(doc_xml)
    body = tree.find(f"{{{W}}}body")

    blocks = []
    for child in body:
        tag = child.tag.split("}")[-1]
        if tag in ("p", "tbl"):
            blocks.append(child)
    _log.debug(f"  _slice_docx: {len(blocks)} blocks, mem={mem()}MB")

    slices = []
    total_slices = (len(blocks) + blocks_per_slice - 1) // blocks_per_slice
    for i in range(0, len(blocks), blocks_per_slice):
        slice_idx = i // blocks_per_slice
        if slice_idx % 10 == 0:
            _log.debug(f"  _slice_docx: slice {slice_idx}/{total_slices}, mem={mem()}MB")
        end = min(i + blocks_per_slice, len(blocks))
        out_path = slices_dir / f"slice_{slice_idx:03d}.docx"

        with zipfile.ZipFile(src, "r") as zin:
            with zipfile.ZipFile(str(out_path), "w", zipfile.ZIP_DEFLATED) as zout:
                doc_xml = zin.read("word/document.xml")
                tree = etree.fromstring(doc_xml)
                body = tree.find(f"{{{W}}}body")
                for child in list(body):
                    body.remove(child)
                for block in blocks[i:end]:
                    body.append(copy.deepcopy(block))
                for item in zin.namelist():
                    if item == "word/document.xml":
                        zout.writestr(item, etree.tostring(tree))
                    else:
                        zout.writestr(item, zin.read(item))

        slices.append({
            "id": f"b{i + 1}-{end}",
            "path": str(out_path),
            "blocks": (i + 1, end),
            "total": len(blocks)
        })
    _log.debug(f"  _slice_docx: done, {len(slices)} slices, mem={mem()}MB")
    return slices


# ─────────────────────────────────────────────────────────────────
# 文字提取
# ─────────────────────────────────────────────────────────────────

def _read_pdf_text(path: str) -> str:
    """PyMuPDF 提取 PDF 文字"""
    import fitz
    _log.debug(f"  _read_pdf_text: opening {path}")
    with fitz.open(path) as doc:
        total = len(doc)
        parts = []
        for i, page in enumerate(doc):
            if i % 20 == 0:
                _log.debug(f"  _read_pdf_text: page {i+1}/{total}")
            text = page.get_text("text")
            if text.strip():
                parts.append(f"--- 第 {i + 1} 页 ---\n{text}")
        _log.debug(f"  _read_pdf_text: done, {len(parts)} pages with text")
        return "\n\n".join(parts)


def _read_docx_text(path: str) -> str:
    """python-docx 提取 DOCX 文字+表格"""
    from docx import Document
    _log.debug(f"  _read_docx_text: opening {path}")
    doc = Document(path)
    parts = []
    total = len(list(doc.element.body))
    _log.debug(f"  _read_docx_text: {total} blocks")

    for idx, elem in enumerate(doc.element.body):
        if idx % 100 == 0:
            _log.debug(f"  _read_docx_text: block {idx}/{total}")
        tag = elem.tag.split("}")[-1]
        if tag == "p":
            text = "".join(
                t.text or "" for t in elem.iter(f"{{{docx_ns}}}t")
            )
            if text.strip():
                parts.append(text)
        elif tag == "tbl":
            tbl_text = _docx_table_to_markdown(elem)
            if tbl_text:
                parts.append(tbl_text)

    _log.debug(f"  _read_docx_text: done, {len(parts)} blocks with content")
    return "\n\n".join(parts)


def _docx_table_to_markdown(tbl_elem) -> str:
    """将 DOCX 表格元素转为 Markdown 表格"""
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    rows = tbl_elem.findall(f".//{{{W}}}tr")
    if not rows:
        return ""

    parts = []
    for ri, tr in enumerate(rows):
        cells = tr.findall(f".//{{{W}}}tc")
        row_text = []
        for tc in cells:
            cell_text = "".join(
                t.text or "" for t in tc.iter(f"{{{W}}}t")
            ).strip()
            row_text.append(cell_text)
        parts.append("| " + " | ".join(row_text) + " |")
        if ri == 0:
            parts.append("| " + " | ".join(["---"] * len(row_text)) + " |")
    return "\n".join(parts)


docx_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _strip_base64_images(text: str) -> str:
    return re.sub(r"!\[([^\]]*)\]\((data:image/[^)]+)\)", "[图片]", text)


# ─────────────────────────────────────────────────────────────────
# 图片提取（统一入口）
# ─────────────────────────────────────────────────────────────────

def _extract_images_from_pdf(pdf_path: str, doc_name: str) -> list[dict]:
    """从 PDF 提取图片，返回图片信息列表（含大小/页码/MD5）"""
    import fitz
    _log.debug(f"  _extract_images_from_pdf: opening {pdf_path}")
    images_dir = _get_images_dir(doc_name)
    index = _load_index(doc_name)
    seen_hashes: set[str] = set(img["md5"] for img in index["images"])
    results = []

    with fitz.open(pdf_path) as doc:
        total_pages = len(doc)
        _log.debug(f"  _extract_images_from_pdf: {total_pages} pages, mem={mem()}MB")
        for page_num in range(len(doc)):
            if page_num % 20 == 0:
                _log.debug(f"  _extract_images_from_pdf: page {page_num+1}/{total_pages}, mem={mem()}MB")
            page = doc[page_num]
            for img_index, img in enumerate(page.get_images(full=True)):
                xref = img[0]
                base_image = doc.extract_image(xref)
                img_bytes = base_image["image"]
                content_hash = hashlib.md5(img_bytes).hexdigest()
                ext = base_image.get("ext", "png").lower()
                if ext not in ("png", "jpg", "jpeg", "gif", "bmp", "webp"):
                    ext = "png"
                img_hash = content_hash[:12]
                out_name = f"{doc_name}_p{page_num + 1}_i{img_index}_{img_hash}.{ext}"
                out_path = images_dir / out_name

                is_new = content_hash not in seen_hashes
                if is_new:
                    seen_hashes.add(content_hash)
                    out_path.write_bytes(img_bytes)
                    index["images"].append({
                        "name": out_name,
                        "md5": content_hash,
                        "path": str(out_path),
                        "page": page_num + 1,
                        "size": len(img_bytes)
                    })

                results.append({
                    "name": out_name,
                    "path": str(out_path),
                    "page": page_num + 1,
                    "size": len(img_bytes),
                    "is_new": is_new
                })

    _log.debug(f"  _extract_images_from_pdf: done, {len(results)} images, mem={mem()}MB")
    _save_index(doc_name, index)
    return results


def _extract_images_from_docx(docx_path: str, doc_name: str) -> list[dict]:
    """从 DOCX 提取图片，返回图片信息列表"""
    _log.debug(f"  _extract_images_from_docx: opening {docx_path}")
    images_dir = _get_images_dir(doc_name)
    index = _load_index(doc_name)
    seen_hashes: set[str] = set(img["md5"] for img in index["images"])
    results = []

    with zipfile.ZipFile(docx_path, "r") as z:
        media_files = [f for f in z.namelist() if f.startswith("word/media/")]
        _log.debug(f"  _extract_images_from_docx: {len(media_files)} media files, mem={mem()}MB")
        for i, mf in enumerate(media_files):
            if i % 50 == 0:
                _log.debug(f"  _extract_images_from_docx: {i}/{len(media_files)}")
            img_bytes = z.read(mf)
            content_hash = hashlib.md5(img_bytes).hexdigest()
            ext = Path(mf).suffix.lower()
            if ext not in (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"):
                ext = ".png"
            img_hash = content_hash[:12]
            out_name = f"{doc_name}_docx_{img_hash}{ext}"
            out_path = images_dir / out_name

            is_new = content_hash not in seen_hashes
            if is_new:
                seen_hashes.add(content_hash)
                out_path.write_bytes(img_bytes)
                index["images"].append({
                    "name": out_name,
                    "md5": content_hash,
                    "path": str(out_path),
                    "size": len(img_bytes)
                })

            results.append({
                "name": out_name,
                "path": str(out_path),
                "size": len(img_bytes),
                "is_new": is_new
            })

    _log.debug(f"  _extract_images_from_docx: done, {len(results)} images, mem={mem()}MB")
    _save_index(doc_name, index)
    return results


# ─────────────────────────────────────────────────────────────────
# 小图 OCR
# ─────────────────────────────────────────────────────────────────

def _ocr_small_image(img_path: str) -> str:
    """对小于阈值的小图进行 OCR，返回识别文字"""
    try:
        from PIL import Image
        import pytesseract
        img = Image.open(img_path)
        text = pytesseract.image_to_string(img, lang="chi_sim+eng")
        return text.strip()
    except Exception:
        return ""


# ─────────────────────────────────────────────────────────────────
# 单文档读取（核心逻辑）
# ─────────────────────────────────────────────────────────────────

def _read_single_document(
    file_path: str,
    doc_name: str,
    mode: str = "auto",
    extract_images: bool = True,
) -> dict:
    """
    读取单个文档（PDF 或 DOCX）
    mode: "pdf_only" | "docx_only" | "auto"
    """
    _log.debug(f"  _read_single: START file={file_path} mode={mode} mem={mem()}MB")
    p = Path(file_path)
    ext = p.suffix.lower()

    slices = []
    text_parts = []
    all_images = []

    if ext == ".pdf" or (mode in ("auto", "pdf_only") and _is_pdf_by_magic(file_path)):
        slices = _slice_pdf(file_path, doc_name)
        _log.debug(f"  _read_single: {len(slices)} PDF slices, mem={mem()}MB")
        for i, sl in enumerate(slices):
            if i % 5 == 0:
                _log.debug(f"  _read_single: PDF slice {i}/{len(slices)}, mem={mem()}MB")
            txt = _read_pdf_text(sl["path"])
            text_parts.append(f"=== [{sl['id']}] ===\n{txt}")
            if extract_images:
                imgs = _extract_images_from_pdf(sl["path"], doc_name)
                all_images.extend(imgs)

    elif ext in (".docx", ".doc") or mode == "docx_only":
        slices = _slice_docx(file_path, doc_name)
        _log.debug(f"  _read_single: {len(slices)} DOCX slices, mem={mem()}MB")
        for i, sl in enumerate(slices):
            if i % 5 == 0:
                _log.debug(f"  _read_single: DOCX slice {i}/{len(slices)}, mem={mem()}MB")
            txt = _read_docx_text(sl["path"])
            text_parts.append(f"=== [{sl['id']}] ===\n{txt}")
            if extract_images:
                imgs = _extract_images_from_docx(sl["path"], doc_name)
                all_images.extend(imgs)

    else:
        txt = _read_generic_text(file_path)
        text_parts.append(txt)

    full_text = "\n\n".join(text_parts)
    _log.debug(f"  _read_single: text done, {len(full_text)} chars, mem={mem()}MB")

    if extract_images:
        _log.debug(f"  _read_single: OCR on {len(all_images)} images")
        ocr_count = 0
        for img in all_images:
            if img["size"] < IMAGE_SIZE_THRESHOLD:
                ocr_text = _ocr_small_image(img["path"])
                if ocr_text:
                    img["ocr"] = ocr_text
                    ocr_count += 1
        _log.debug(f"  _read_single: OCR done, {ocr_count} succeeded, mem={mem()}MB")

    written = False
    content_path = None
    if len(full_text) > MAX_CHARS_RETURN:
        content_path = _get_doc_dir(doc_name) / "content.md"
        content_path.write_text(full_text, encoding="utf-8")
        written = True

    _log.debug(f"  _read_single: DONE mem={mem()}MB")
    return {
        "text": full_text,
        "images": all_images,
        "slices": slices,
        "paired": False,
        "paired_file": None,
        "written": written,
        "content_path": str(content_path) if written else None
    }


# ─────────────────────────────────────────────────────────────────
# 配对文档读取
# ─────────────────────────────────────────────────────────────────

def _read_paired_documents(
    pdf_path: str, docx_path: str, extract_images: bool = True
) -> dict:
    pdf_name = Path(pdf_path).stem
    doc_name = pdf_name
    _log.debug(f"  _read_paired: START pdf={pdf_name} mem={mem()}MB")

    _log.debug(f"  _read_paired: slicing PDF")
    pdf_slices = _slice_pdf(pdf_path, doc_name)
    _log.debug(f"  _read_paired: {len(pdf_slices)} PDF slices, mem={mem()}MB")

    all_images = []
    if extract_images:
        _log.debug(f"  _read_paired: extracting images from {len(pdf_slices)} slices")
        for i, sl in enumerate(pdf_slices):
            if i % 5 == 0:
                _log.debug(f"  _read_paired: image slice {i}/{len(pdf_slices)}, mem={mem()}MB")
            imgs = _extract_images_from_pdf(sl["path"], doc_name)
            all_images.extend(imgs)
        _log.debug(f"  _read_paired: images done, {len(all_images)} total, mem={mem()}MB")

    _log.debug(f"  _read_paired: slicing DOCX")
    docx_slices = _slice_docx(docx_path, doc_name)
    _log.debug(f"  _read_paired: {len(docx_slices)} DOCX slices, mem={mem()}MB")

    text_parts = []
    for i, sl in enumerate(docx_slices):
        if i % 5 == 0:
            _log.debug(f"  _read_paired: text slice {i}/{len(docx_slices)}, mem={mem()}MB")
        txt = _read_docx_text(sl["path"])
        text_parts.append(f"=== [{sl['id']}] ===\n{txt}")

    full_text = "\n\n".join(text_parts)
    _log.debug(f"  _read_paired: text done, {len(full_text)} chars, mem={mem()}MB")

    if extract_images:
        _log.debug(f"  _read_paired: OCR on {len(all_images)} images")
        ocr_count = 0
        for img in all_images:
            if img["size"] < IMAGE_SIZE_THRESHOLD:
                ocr_text = _ocr_small_image(img["path"])
                if ocr_text:
                    img["ocr"] = ocr_text
                    ocr_count += 1
        _log.debug(f"  _read_paired: OCR done, {ocr_count} succeeded, mem={mem()}MB")

    written = False
    content_path = None
    if len(full_text) > MAX_CHARS_RETURN:
        content_path = _get_doc_dir(doc_name) / "content.md"
        content_path.write_text(full_text, encoding="utf-8")
        written = True
        _log.debug(f"  _read_paired: content written to {content_path}")

    _log.debug(f"  _read_paired: DONE mem={mem()}MB")
    return {
        "text": full_text,
        "images": all_images,
        "slices": docx_slices,
        "paired": True,
        "paired_file": docx_path,
        "written": written,
        "content_path": str(content_path) if written else None
    }


# ─────────────────────────────────────────────────────────────────
# 辅助
# ─────────────────────────────────────────────────────────────────

def _is_pdf_by_magic(path: str) -> bool:
    try:
        with open(path, "rb") as f:
            header = f.read(8)
        return header.startswith(b"%PDF")
    except Exception:
        return False


def _read_generic_text(path: str) -> str:
    for enc in ("utf-8", "gbk", "gb2312", "utf-16"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    size = os.path.getsize(path)
    return f"# 二进制文件 (无法解码)\n大小: {size:,} 字节"


# ─────────────────────────────────────────────────────────────────
# MCP 工具
# ─────────────────────────────────────────────────────────────────

@mcp.tool()
def read_document(file_path: str, fast: bool = False) -> str:
    _log.debug(f"read_document CALLED file={file_path} fast={fast} mem={mem()}MB")
    if not os.path.isfile(file_path):
        return f"[错误] 文件不存在: {file_path}"

    doc_name = Path(file_path).stem.replace(" ", "_")[:50]
    paired = _find_paired_file(file_path)
    extract_imgs = not fast

    try:
        if paired is None:
            result = _read_single_document(file_path, doc_name, extract_images=extract_imgs)
        else:
            ext = Path(file_path).suffix.lower()
            if ext == ".pdf":
                result = _read_paired_documents(file_path, paired, extract_images=extract_imgs)
            else:
                result = _read_paired_documents(paired, file_path, extract_images=extract_imgs)
    except Exception:
        _log.exception(f"  read_document CRASHED: {traceback.format_exc()}")
        return f"[错误] 处理失败:\n{traceback.format_exc()}"

    # 构造返回
    img_lines = []
    ocr_count = 0
    for img in result["images"]:
        size_kb = img["size"] // 1024
        ocr_hint = f" [OCR: {img.get('ocr', '')}]" if img.get('ocr') else ""
        img_lines.append(f"[图片: file://{img['path']}] ({size_kb}KB){ocr_hint}")
        if img.get('ocr'):
            ocr_count += 1

    slice_info = ", ".join(s["id"] for s in result["slices"])
    img_summary = f"{len(result['images'])}张图片"
    if ocr_count > 0:
        img_summary += f"（含{ocr_count}张已OCR）"

    paired_line = (
        f"配对: {'是 → ' + Path(result['paired_file']).name if result['paired'] else '未找到配对文件'}\n"
    )
    if not result["paired"]:
        ext = Path(file_path).suffix.lower()
        if ext == ".pdf":
            paired_line += (
                "⚠️ 只提供 PDF: 文字可能乱码（CID-font），图片可正常提取\n"
                "提示: 如有 DOCX 版本，建议 read_document_pair(pdf, docx)\n"
                + ("⚠️ fast=True 已跳过图片提取，图片请用 extract_images() 单独提取\n" if fast else "      继续干活...\n")
            )
        else:
            paired_line += (
                "⚠️ 只提供 DOCX: 文字表格可正常提取，可能无法提取图片\n"
                "提示: 如有 PDF 版本，建议 read_document_pair(pdf, docx)\n"
                + ("⚠️ fast=True 已跳过图片提取，图片请用 extract_images() 单独提取\n" if fast else "      继续干活...\n")
            )

    # 结构化状态 JSON（agent 可直接解析）
    import json
    status = {
        "doc": Path(file_path).name,
        "paired": result["paired"],
        "paired_file": Path(result["paired_file"]).name if result["paired_file"] else None,
        "need_pairing": not result["paired"],
        "fast": fast,
        "images_extracted": extract_imgs,
        "slice_count": len(result["slices"]),
        "slices": [s["id"] for s in result["slices"]],
        "image_count": len(result["images"]),
        "ocr_count": ocr_count,
        "content_written": result["written"],
        "content_path": result["content_path"],
        "chars": len(result["text"]),
        "mode": "paired" if result["paired"] else ("pdf_single" if (Path(file_path).suffix.lower() == ".pdf") else "docx_single"),
    }
    status_json = json.dumps(status, ensure_ascii=False)

    header = (
        f"[STATUS]\n{status_json}\n[/STATUS]\n\n"
        f"# {Path(file_path).name}\n"
        + paired_line
        + f"分片: {len(result['slices'])}片 ({slice_info})\n"
        + f"图片: {img_summary}\n"
    )

    if result["written"]:
        footer = (
            f"\n---\n[内容已写入文件: {result['content_path']}]\n"
            f"[完整内容请读取该文件]"
        )
    else:
        footer = f"\n--- 字符: {len(result['text'])} ---\n"

    if img_lines:
        footer += "\n" + "\n".join(img_lines)

    return header + result["text"][:MAX_CHARS_RETURN] + footer


@mcp.tool()
def read_document_pair(pdf_path: str, docx_path: str) -> str:
    """显式配对读取：PDF 提取图片，DOCX 提取文字和表格"""
    _log.debug(f"read_document_pair CALLED pdf={pdf_path} docx={docx_path} mem={mem()}MB")
    if not os.path.isfile(pdf_path):
        return f"[错误] PDF 不存在: {pdf_path}"
    if not os.path.isfile(docx_path):
        return f"[错误] DOCX 不存在: {docx_path}"

    doc_name = Path(pdf_path).stem.replace(" ", "_")[:50]
    try:
        result = _read_paired_documents(pdf_path, docx_path)
    except Exception:
        _log.exception(f"  read_document_pair CRASHED: {traceback.format_exc()}")
        return f"[错误] 处理失败:\n{traceback.format_exc()}"

    img_lines = []
    ocr_count = 0
    for img in result["images"]:
        size_kb = img["size"] // 1024
        ocr_hint = f" [OCR: {img.get('ocr', '')}]" if img.get('ocr') else ""
        img_lines.append(f"[图片: file://{img['path']}] ({size_kb}KB){ocr_hint}")
        if img.get('ocr'):
            ocr_count += 1

    slice_info = ", ".join(s["id"] for s in result["slices"])
    img_summary = f"{len(result['images'])}张图片"
    if ocr_count > 0:
        img_summary += f"（含{ocr_count}张已OCR）"

    header = (
        f"# {Path(pdf_path).name} + {Path(docx_path).name}\n"
        f"配对: 是（PDF图片 + DOCX文字表格）\n"
        f"分片: {len(result['slices'])}片 ({slice_info})\n"
        f"图片: {img_summary}\n"
    )

    if result["written"]:
        footer = (
            f"\n---\n[内容已写入文件: {result['content_path']}]\n"
            f"[完整内容请读取该文件]"
        )
    else:
        footer = f"\n--- 字符: {len(result['text'])} ---\n"

    if img_lines:
        footer += "\n" + "\n".join(img_lines)

    return header + result["text"][:MAX_CHARS_RETURN] + footer


@mcp.tool()
def extract_images(file_path: str, page_range: str = "") -> str:
    """
    专门提取图片。返回图片列表。
    page_range: 可选，如 "1-5" 或 "1,3,5"
    """
    if not os.path.isfile(file_path):
        return f"[错误] 文件不存在: {file_path}"

    doc_name = Path(file_path).stem.replace(" ", "_")[:50]
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf" or _is_pdf_by_magic(file_path):
        if page_range:
            import fitz
            with fitz.open(file_path) as doc:
                indices = []
                for part in page_range.split(","):
                    part = part.strip()
                    if "-" in part:
                        s, e = part.split("-", 1)
                        indices.extend(range(int(s) - 1, int(e)))
                    else:
                        indices.append(int(part) - 1)
                with fitz.open() as tmp:
                    for idx in indices:
                        if 0 <= idx < len(doc):
                            tmp.insert_pdf(doc, from_page=idx, to_page=idx)
                    tmp_path = tempfile.mktemp(suffix=".pdf")
                    tmp.save(tmp_path)
            imgs = _extract_images_from_pdf(tmp_path, doc_name + "_range")
            os.unlink(tmp_path)
        else:
            imgs = _extract_images_from_pdf(file_path, doc_name)
    else:
        imgs = _extract_images_from_docx(file_path, doc_name)

    lines = [f"共 {len(imgs)} 张图片:\n"]
    for img in imgs:
        size_kb = img["size"] // 1024
        ocr_hint = f" [OCR] {img.get('ocr', '')}" if img.get('ocr') else ""
        lines.append(f"[图片: file://{img['path']}] ({size_kb}KB){ocr_hint}")

    return "\n".join(lines)


@mcp.tool()
def ocr_image(image_path: str) -> str:
    """
    对单张图片进行 OCR 识别（内置 tesseract）。
    支持 PNG/JPG/BMP 等常见格式。
    返回识别文字，失败返回空字符串。
    """
    if not os.path.isfile(image_path):
        return f"[错误] 文件不存在: {image_path}"
    _log.debug(f"ocr_image CALLED path={image_path} mem={mem()}MB")
    text = _ocr_small_image(image_path)
    if text:
        _log.debug(f"ocr_image: OK, {len(text)} chars")
    else:
        _log.debug(f"ocr_image: no text found")
    return text


@mcp.tool()
def slice_document(file_path: str, pages_per_slice: int = 5) -> str:
    """
    分片工具：只做分片，返回分片信息（不读取内容）。
    用于提前了解分片结构，或单独控制分片逻辑。
    返回: JSON 格式分片列表
    """
    import json
    _log.debug(f"slice_document CALLED file={file_path} mem={mem()}MB")
    if not os.path.isfile(file_path):
        return f'[{{"error": "文件不存在: {file_path}"}}]'

    p = Path(file_path)
    ext = p.suffix.lower()
    doc_name = p.stem.replace(" ", "_")[:50]

    try:
        if ext == ".pdf" or _is_pdf_by_magic(file_path):
            import fitz
            doc = fitz.open(file_path)
            total = len(doc)
            slices = []
            for i in range(0, total, pages_per_slice):
                end = min(i + pages_per_slice, total)
                slices.append({
                    "slice_id": f"p{i+1}-{end}",
                    "page_range": f"{i+1}-{end}",
                    "page_count": end - i,
                    "slice_index": i // pages_per_slice,
                })
            return json.dumps({
                "doc_name": doc_name,
                "total_pages": total,
                "slice_count": len(slices),
                "slices": slices,
                "paired_file": _find_paired_file(file_path),
            }, ensure_ascii=False, indent=2)

        elif ext in (".docx", ".doc"):
            from docx import Document
            doc = Document(file_path)
            blocks = sum(1 for _ in doc.element.body)
            slices = []
            for i in range(0, blocks, SLICE_BLOCKS):
                end = min(i + SLICE_BLOCKS, blocks)
                slices.append({
                    "slice_id": f"b{i+1}-{end}",
                    "block_range": f"{i+1}-{end}",
                    "block_count": end - i,
                    "slice_index": i // SLICE_BLOCKS,
                })
            return json.dumps({
                "doc_name": doc_name,
                "total_blocks": blocks,
                "slice_count": len(slices),
                "slices": slices,
                "paired_file": _find_paired_file(file_path),
            }, ensure_ascii=False, indent=2)

        else:
            return f'[{{"error": "不支持分片格式: {ext}"}}]'

    except Exception as e:
        _log.exception(f"  slice_document CRASHED for {file_path}: {traceback.format_exc()}")
        return f'[{{"error": "{type(e).__name__}: {e}"}}]'


@mcp.tool()
def get_document_info(file_path: str) -> str:
    """返回文档元信息：页数、分片数、图片数、配对建议"""
    _log.debug(f"get_document_info CALLED file={file_path} mem={mem()}MB")
    if not os.path.isfile(file_path):
        return f"[错误] 文件不存在: {file_path}"

    p = Path(file_path)
    size = os.path.getsize(file_path)
    doc_name = p.stem.replace(" ", "_")[:50]
    ext = p.suffix.lower()

    info = [
        f"文件: {p.name}",
        f"大小: {size / 1024:.0f} KB ({size / 1024**2:.1f} MB)",
    ]

    paired = _find_paired_file(file_path)
    if paired:
        info.append(f"配对文件: 有 → {Path(paired).name}")
        info.append("建议: 使用 read_document_pair() 显式配对读取")
    else:
        info.append(f"配对文件: 无")
        info.append("提示: 如有另一格式版本，建议提供以获得更完整结果")

    if ext == ".pdf" or _is_pdf_by_magic(file_path):
        try:
            import fitz
            with fitz.open(file_path) as doc:
                total = len(doc)
                slices = (total + SLICE_PAGES - 1) // SLICE_PAGES
                info.append(f"页数: {total}")
                info.append(f"预估分片数: {slices}片 (每{SLICE_PAGES}页)")
        except Exception as e:
            _log.warning(f"  get_document_info: fitz failed for {file_path}: {e}")
    elif ext in (".docx", ".doc"):
        try:
            from docx import Document
            doc = Document(file_path)
            blocks = sum(1 for _ in doc.element.body)
            slices = (blocks + SLICE_BLOCKS - 1) // SLICE_BLOCKS
            info.append(f"段落: {len(doc.paragraphs)}")
            info.append(f"表格: {len(doc.tables)}")
            info.append(f"预估分片数: {slices}片 (每{SLICE_BLOCKS}block)")
        except Exception as e:
            _log.warning(f"  get_document_info: docx failed for {file_path}: {e}")

    # 检查图片池
    index = _load_index(doc_name)
    if index["images"]:
        info.append(f"图片池已有: {len(index['images'])} 张")

    return "\n".join(info)


@mcp.tool()
def list_supported_files() -> str:
    return (
        "支持的文件格式:\n"
        "  PDF:  .pdf\n"
        "  Word: .docx, .doc\n"
        "  文本: .txt, .md, .json, .xml, .csv, .html\n"
        "\n"
        "工具接口:\n"
        "  read_document(path)         - 主工具，自动配对检测\n"
        "  read_document_pair(pdf, docx) - 显式配对读取\n"
        "  extract_images(path)        - 专门提取图片\n"
        "  get_document_info(path)      - 文档元信息\n"
    )


# ─────────────────────────────────────────────────────────────────
# 启动
# ─────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import resource

    def mem():
        r = resource.getrusage(resource.RUSAGE_SELF)
        return r.ru_maxrss // 1024

    _log.debug(f"PID={os.getpid()} STARTING mem={mem()}MB")
    BASE_DIR.mkdir(parents=True, exist_ok=True)
    _log.debug(f"BASE_DIR ready mem={mem()}MB")
    _log.debug(f"calling mcp.run() mem={mem()}MB")
    mcp.run()
    _log.debug(f"mcp.run() returned")
